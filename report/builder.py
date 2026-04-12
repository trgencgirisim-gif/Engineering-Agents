"""report/builder.py — DOCX and PDF report generation entry points.

Phase 3.6: generate_pdf_report() produces true PDF output via ReportLab
(if available) instead of being an alias for the DOCX generator.
"""

import io
import datetime
from typing import List, Dict

from docx import Document
from docx.shared import Pt, Cm

from report.styles import (
    SZ_ABSTRACT, SZ_SMALL,
    C_BLACK, C_NAVY, C_GREY,
    _run, _set_line_spacing, _set_margins, _setup_styles,
    _h1, _body, _extract_abstract,
)
from report.sections import (
    _build_cover, _build_introduction, _build_methodology,
    _build_findings, _build_discussion, _build_conclusions,
    _build_references, _build_appendix_a, _build_appendix_b,
    _setup_running_headers,
)


def generate_docx_report(
    brief:        str,
    final_report: str,
    domains:      List[str],
    round_scores: List[Dict] = None,
    agent_log:    List[Dict] = None,
    total_cost:   float      = 0.0,
    kur:          float      = 44.0,
    mode:         int        = 4,
    max_rounds:   int        = 3,
) -> bytes:

    round_scores = round_scores or []
    agent_log    = agent_log    or []

    doc_no = f"EAI-{datetime.datetime.now().strftime('%Y%m%d-%H%M')}-M{mode}"

    mode_labels = {1: "Single Agent", 2: "Dual Agent",
                   3: "Semi-Automatic", 4: "Full Automatic"}
    mode_label  = mode_labels.get(mode, "Custom")
    brief_short = brief[:90].replace("\n", " ").strip()

    # Document setup
    doc = Document()
    _set_margins(doc, top=2.5, right=2.0, bottom=2.5, left=2.5)
    _setup_styles(doc)

    # 1. Cover page
    _build_cover(doc, brief, domains, mode, total_cost, kur,
                 round_scores, agent_log, mode_label, doc_no)

    # 2. Abstract
    _h1(doc, "", "Abstract")
    abstract = _extract_abstract(final_report)
    if abstract:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.8)
        p.paragraph_format.right_indent = Cm(0.8)
        p.paragraph_format.space_after  = Pt(8)
        _set_line_spacing(p, 1.15)
        _run(p, abstract, italic=True, size=SZ_ABSTRACT, color=C_BLACK)
    else:
        _body(doc, f"This report presents the results of a multi-agent engineering "
              f"analysis of the following problem: {brief_short}. "
              f"The analysis was conducted in {mode_label} and covers "
              f"{len(domains)} engineering domain(s): {', '.join(domains)}.")

    # Keyword line
    kw = doc.add_paragraph()
    kw.paragraph_format.left_indent = Cm(0.8)
    kw.paragraph_format.space_after = Pt(0)
    _run(kw, "Keywords:  ", bold=True, size=SZ_SMALL, color=C_NAVY)
    _run(kw, ", ".join(domains) + f"; multi-agent analysis; engineering AI; {mode_label.lower()}",
         italic=True, size=SZ_SMALL, color=C_GREY)

    doc.add_page_break()

    # 3. Introduction
    _build_introduction(doc, brief, domains, mode)

    # 4. Methodology
    _build_methodology(doc, mode, domains, max_rounds, round_scores)

    # 5. Technical Findings
    _build_findings(doc, final_report, domains)

    # 6. Discussion
    _build_discussion(doc, final_report, round_scores)

    # 7. Conclusions & Recommendations
    _build_conclusions(doc, final_report)

    doc.add_page_break()

    # 8. References
    _build_references(doc, final_report)

    doc.add_page_break()

    # Appendix A
    _build_appendix_a(doc, domains, mode, round_scores,
                      total_cost, kur, agent_log, max_rounds)

    doc.add_page_break()

    # Appendix B
    _build_appendix_b(doc, agent_log)

    # Running headers / footers
    _setup_running_headers(doc, brief_short, doc_no)

    # Serialize
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def generate_pdf_report(
    brief:        str,
    final_report: str,
    domains:      List[str],
    round_scores: List[Dict] = None,
    agent_log:    List[Dict] = None,
    total_cost:   float      = 0.0,
    kur:          float      = 44.0,
    mode:         int        = 4,
    max_rounds:   int        = 3,
) -> bytes:
    """
    Phase 3.6: Generate a true PDF report using ReportLab.

    Falls back to DOCX bytes if ReportLab is not installed, so existing
    callers are never broken.
    """
    try:
        return _generate_pdf_reportlab(
            brief=brief,
            final_report=final_report,
            domains=domains,
            round_scores=round_scores or [],
            agent_log=agent_log or [],
            total_cost=total_cost,
            kur=kur,
            mode=mode,
            max_rounds=max_rounds,
        )
    except ImportError:
        # ReportLab not installed — fall back to DOCX
        return generate_docx_report(
            brief=brief, final_report=final_report, domains=domains,
            round_scores=round_scores, agent_log=agent_log,
            total_cost=total_cost, kur=kur, mode=mode, max_rounds=max_rounds,
        )


def _generate_pdf_reportlab(
    brief: str,
    final_report: str,
    domains: List[str],
    round_scores: List[Dict],
    agent_log: List[Dict],
    total_cost: float,
    kur: float,
    mode: int,
    max_rounds: int,
) -> bytes:
    """Internal ReportLab-based PDF builder."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        HRFlowable, PageBreak,
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

    buf = io.BytesIO()
    doc_no = f"EAI-{datetime.datetime.now().strftime('%Y%m%d-%H%M')}-M{mode}"
    mode_labels = {1: "Single Agent", 2: "Dual Agent",
                   3: "Semi-Automatic", 4: "Full Automatic"}
    mode_label = mode_labels.get(mode, "Custom")
    brief_short = brief[:120].replace("\n", " ").strip()

    # ── Page setup ───────────────────────────────────────────
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2.5 * cm, rightMargin=2.0 * cm,
        topMargin=2.5 * cm, bottomMargin=2.5 * cm,
        title=f"Engineering Analysis — {brief_short}",
        author="Engineering AI Multi-Agent System",
    )

    # ── Styles ───────────────────────────────────────────────
    styles = getSampleStyleSheet()
    NAVY = colors.HexColor("#1a2747")
    GREY = colors.HexColor("#555555")

    style_h1 = ParagraphStyle("H1", parent=styles["Heading1"],
                               textColor=NAVY, fontSize=14, spaceAfter=8)
    style_h2 = ParagraphStyle("H2", parent=styles["Heading2"],
                               textColor=NAVY, fontSize=12, spaceAfter=6)
    style_body = ParagraphStyle("Body", parent=styles["Normal"],
                                 fontSize=10, leading=14, alignment=TA_JUSTIFY,
                                 spaceAfter=6)
    style_mono = ParagraphStyle("Mono", parent=styles["Normal"],
                                 fontName="Courier", fontSize=9, leading=12,
                                 textColor=GREY, spaceAfter=4)
    style_center = ParagraphStyle("Center", parent=styles["Normal"],
                                   alignment=TA_CENTER, fontSize=10)
    style_title = ParagraphStyle("Title", parent=styles["Title"],
                                  textColor=NAVY, fontSize=20, spaceAfter=12,
                                  alignment=TA_CENTER)

    story = []

    def h1(text):
        story.append(Paragraph(text, style_h1))
        story.append(HRFlowable(width="100%", thickness=1, color=NAVY, spaceAfter=6))

    def h2(text):
        story.append(Paragraph(text, style_h2))

    def body(text):
        # Escape HTML-special chars for ReportLab
        text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(text, style_body))

    def mono(text):
        text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(text, style_mono))

    def sp(n=1):
        story.append(Spacer(1, n * 0.3 * cm))

    # ── Cover ────────────────────────────────────────────────
    story.append(Spacer(1, 2 * cm))
    story.append(Paragraph("Engineering Analysis Report", style_title))
    sp(2)
    story.append(Paragraph(brief_short, style_center))
    sp()
    story.append(Paragraph(
        f"Document: {doc_no}  |  Mode: {mode_label}  |  "
        f"Domains: {', '.join(domains[:4])}" +
        (f" +{len(domains)-4} more" if len(domains) > 4 else ""),
        style_center
    ))
    sp()
    story.append(Paragraph(
        f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}  |  "
        f"Cost: ${total_cost:.4f} USD ({total_cost * kur:.2f} TRY)",
        style_center
    ))
    story.append(PageBreak())

    # ── Quality progression table ─────────────────────────────
    if round_scores:
        h1("Quality Progression")
        sp()
        tdata = [["Round", "Score", "Notes"]]
        for r in round_scores:
            rnd = r.get("round", "?")
            score = r.get("score", "-")
            notes = "; ".join(r.get("key_changes", [])[:2])[:60]
            tdata.append([str(rnd), str(score), notes])
        t = Table(tdata, colWidths=[2 * cm, 3 * cm, None])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), NAVY),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f5f5")]),
        ]))
        story.append(t)
        sp()

    # ── Final report body ─────────────────────────────────────
    h1("Technical Findings & Analysis")
    sp()
    for line in final_report.splitlines():
        stripped = line.strip()
        if not stripped:
            sp(0.5)
        elif stripped.startswith("## "):
            h2(stripped[3:])
        elif stripped.startswith("# "):
            h1(stripped[2:])
        elif stripped.startswith("```") or stripped.startswith("    "):
            mono(stripped.lstrip("`").lstrip())
        else:
            body(stripped)

    story.append(PageBreak())

    # ── Agent cost appendix ────────────────────────────────────
    if agent_log:
        h1("Appendix — Agent Run Summary")
        sp()
        tdata = [["Agent", "Cost (USD)", "Input T.", "Output T.", "Score"]]
        for entry in agent_log[:50]:
            tdata.append([
                str(entry.get("ajan", entry.get("agent", "")))[:30],
                f"${entry.get('maliyet', 0):.5f}",
                str(entry.get("giris", entry.get("input_tokens", "-"))),
                str(entry.get("cikis", entry.get("output_tokens", "-"))),
                str(entry.get("puan", "-")),
            ])
        t = Table(tdata, colWidths=[6 * cm, 2.5 * cm, 2.5 * cm, 2.5 * cm, 2 * cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), NAVY),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f5f5")]),
        ]))
        story.append(t)

    # ── Build ─────────────────────────────────────────────────
    doc.build(story)
    buf.seek(0)
    return buf.read()
