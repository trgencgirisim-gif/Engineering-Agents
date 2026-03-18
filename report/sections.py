"""report/sections.py — Document section builders (cover, intro, methodology, findings, etc.)."""

import re
import datetime
from typing import List, Dict

from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from report.styles import (
    F, SZ_BODY, SZ_H1, SZ_H2, SZ_SMALL, SZ_ABSTRACT,
    C_BLACK, C_NAVY, C_NAVY_HX, C_GREY, C_MUTED, C_OK, C_WARN, C_ERR,
    _run, _set_line_spacing, _shade, _no_border_table,
    _h1, _h2, _h3, _h_rule, _appendix_h1,
    _body, _caption, _note_box, _make_table,
    _render_body, _strip_markup, _extract_abstract, _extract_section, _parse_sections,
)
from report.charts import (
    chart_quality_scores, chart_agent_cost, chart_rpn, _embed_chart,
)


# ═══════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════

def _build_cover(doc, brief: str, domains: list, mode: int,
                 total_cost: float, kur: float, round_scores: list,
                 agent_log: list, mode_label: str, doc_no: str):
    """Academic-style cover page."""
    for _ in range(5):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    sl = doc.add_paragraph()
    sl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(sl, "ENGINEERING AI  \u00b7  MULTI-AGENT ANALYSIS SYSTEM",
         size=8, color=C_MUTED)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

    dt = doc.add_paragraph()
    dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(dt, "TECHNICAL ANALYSIS REPORT", bold=True, size=20, color=C_NAVY)
    dt.paragraph_format.space_after = Pt(4)

    _h_rule(doc, C_NAVY_HX, 14)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(14)

    brief_short = brief[:160].replace("\n", " ").strip()
    if len(brief) > 160:
        brief_short += "..."

    pt_label = doc.add_paragraph()
    pt_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(pt_label, "ANALYSIS SUBJECT", size=8, color=C_MUTED)
    pt_label.paragraph_format.space_after = Pt(3)

    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(pt, brief_short, size=11, color=C_BLACK)
    _set_line_spacing(pt, 1.2)
    pt.paragraph_format.space_after = Pt(20)

    mode_labels = {1: "Mode 1 \u2014 Single Agent", 2: "Mode 2 \u2014 Dual Agent",
                   3: "Mode 3 \u2014 Semi-Automatic", 4: "Mode 4 \u2014 Full Automatic"}
    best_score  = max((r.get("puan", 0) for r in (round_scores or [])), default=None)
    n_domains   = len(domains)
    n_agents    = len(agent_log) if agent_log else 0
    n_rounds    = len(round_scores) if round_scores else 1

    meta_rows = [
        ("Document No.",     doc_no),
        ("Date",             datetime.datetime.now().strftime("%B %d, %Y")),
        ("Analysis Mode",    mode_labels.get(mode, mode_label)),
        ("Engineering Domains", ", ".join(domains) if domains else "\u2014"),
        ("Analysis Rounds",  str(n_rounds)),
        ("Total Agents",     str(n_agents)),
        ("Quality Score",    f"{best_score}/100" if best_score else "\u2014"),
        ("Total API Cost",   f"${total_cost:.4f} USD   (\u2248 {total_cost * kur:.2f} TL)"),
    ]

    mt = doc.add_table(rows=len(meta_rows), cols=2)
    _no_border_table(mt)
    mt.columns[0].width = Cm(4.5)
    mt.columns[1].width = Cm(11.5)

    for ri, (label, value) in enumerate(meta_rows):
        bg = "F5F7FA" if ri % 2 == 0 else "FAFBFC"
        lc = mt.rows[ri].cells[0]
        vc = mt.rows[ri].cells[1]
        for cell, txt, bold in ((lc, label, True), (vc, value, False)):
            _shade(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.left_indent  = Cm(0.3)
            _run(p, txt, bold=bold, size=9.5,
                 color=C_NAVY if bold else C_BLACK)

    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(10)
    _h_rule(doc, C_NAVY_HX, 8)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(8)

    fn = doc.add_paragraph()
    fn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(fn, "This document was generated automatically by the Engineering AI multi-agent analysis system. "
         "Results should be reviewed by a qualified engineer before use in design decisions.",
         italic=True, size=8.5, color=C_MUTED)

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# SECTION 1 — INTRODUCTION
# ═══════════════════════════════════════════════════════════════

def _build_introduction(doc, brief: str, domains: list, mode: int,
                        assumptions: list = None):
    _h1(doc, "1.", "Introduction")

    _h2(doc, "1.1", "Problem Statement")
    _body(doc, brief.strip())

    _h2(doc, "1.2", "Scope and Objectives")
    scope_text = (
        f"This analysis addresses the engineering problem stated above by deploying "
        f"{len(domains)} specialist domain agent(s): {', '.join(domains)}. "
        f"The analysis was conducted in accordance with Mode {mode} of the "
        f"Engineering AI multi-agent framework. "
        f"Each domain agent provides independent analysis from either a theoretical "
        f"(Expert A) or applied (Expert B) perspective, followed by cross-validation, "
        f"observer evaluation, and synthesis into this report."
    )
    _body(doc, scope_text)

    _h2(doc, "1.3", "Assumptions and Limitations")
    if assumptions:
        for a in assumptions:
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Cm(0.6)
            p.paragraph_format.first_line_indent = Cm(-0.4)
            _run(p, "\u2022  ", bold=True, size=SZ_BODY, color=C_NAVY)
            _run(p, a, size=SZ_BODY, color=C_BLACK)
    else:
        _body(doc, (
            "Assumptions embedded within the analysis are explicitly labeled "
            "[ASSUMPTION] within the Technical Findings section. Conservative "
            "engineering estimates have been applied where specific input data "
            "was unavailable, and each such estimate is documented with its "
            "classification (standard simplification, problem-specific, or "
            "conservative bound) and its quantitative impact on the result. "
            "This analysis represents a preliminary multi-perspective assessment "
            "and does not substitute for detailed design calculations, "
            "site-specific measurements, or physical testing."
        ))


# ═══════════════════════════════════════════════════════════════
# SECTION 2 — METHODOLOGY
# ═══════════════════════════════════════════════════════════════

def _build_methodology(doc, mode: int, domains: list, max_rounds: int,
                       round_scores: list):
    _h1(doc, "2.", "Methodology")

    _h2(doc, "2.1", "Analysis Framework")
    mode_descs = {
        1: {
            "label": "Single-Agent Analysis (Mode 1)",
            "body": (
                f"The analysis employed Mode 1 (Single-Agent), deploying one specialist "
                f"expert per engineering domain. Each domain agent performed independent "
                f"analysis, producing quantitative findings, safety assessments, and "
                f"domain-specific recommendations. This mode is optimised for rapid "
                f"preliminary assessment where a single technical perspective per domain "
                f"is sufficient. A total of {len(domains)} domain(s) were activated: "
                f"{', '.join(domains)}."
            )
        },
        2: {
            "label": "Dual-Agent Analysis (Mode 2)",
            "body": (
                f"The analysis employed Mode 2 (Dual-Agent), deploying two specialist "
                f"experts per engineering domain \u2014 Expert A providing theoretical "
                f"analysis and Expert B providing applied, field-validated perspective. "
                f"Positions were compared and conflicts resolved by a dedicated "
                f"Conflict Resolution agent. This mode provides a theory-versus-practice "
                f"dialogue that surfaces design tensions not visible from a single "
                f"perspective. Active domains: {', '.join(domains)}."
            )
        },
        3: {
            "label": "Semi-Automatic Full-Loop Analysis (Mode 3)",
            "body": (
                f"The analysis employed Mode 3 (Semi-Automatic), combining user-supplied "
                f"clarifications with a full multi-round feedback loop. The Prompt "
                f"Engineer agent first identified missing parameters and presented them "
                f"to the user; responses were used to strengthen the problem brief "
                f"before domain analysis commenced. Dual-perspective domain agents "
                f"then executed across {len(domains)} domain(s) over up to "
                f"{max_rounds} rounds, with Observer-directed iteration. "
                f"Active domains: {', '.join(domains)}."
            )
        },
        4: {
            "label": "Full-Automatic Deep Analysis (Mode 4)",
            "body": (
                f"The analysis employed Mode 4 (Full-Automatic), the highest-fidelity "
                f"mode. The Prompt Engineer agent autonomously identified and resolved "
                f"missing parameters before deploying dual-perspective domain agents "
                f"across {len(domains)} domain(s). Up to {max_rounds} iterative "
                f"analysis rounds were executed, with the Observer meta-agent providing "
                f"quality-directed feedback between rounds. Post-loop synthesis, "
                f"alternative scenario, cost, standards, and simulation agents "
                f"supplemented the domain findings. Active domains: "
                f"{', '.join(domains)}."
            )
        },
    }
    m_info = mode_descs.get(mode, {
        "label": f"Mode {mode}",
        "body": f"Analysis mode {mode} was applied across {len(domains)} domain(s): {', '.join(domains)}."
    })
    _body(doc, m_info["body"])

    _h2(doc, "2.2", "Quality Assurance")
    if round_scores:
        score_str = ", ".join(f"Round {r['tur']}: {r.get('puan','\u2014')}/100"
                              for r in round_scores)
        _body(doc, (
            f"Each analysis round was evaluated by an independent Observer meta-agent "
            f"using a weighted rubric: technical accuracy (30%), internal consistency "
            f"(25%), assumption transparency (20%), analysis depth (15%), and "
            f"cross-validation quality (10%). Scores this analysis: {score_str}. "
            f"A score of 85/100 or above triggers early termination."
        ))
    else:
        _body(doc, (
            "Quality control was performed by an independent Observer meta-agent "
            "evaluating technical accuracy, internal consistency, assumption transparency, "
            "analysis depth, and cross-agent consistency."
        ))

    _h2(doc, "2.3", "Domain Coverage")
    rows = [[str(i+1), d] for i, d in enumerate(domains)]
    _make_table(doc,
                ["No.", "Engineering Domain"],
                rows,
                [1.0, 13.0],
                caption="Table 2.1.  Engineering domains active in this analysis.")


# ═══════════════════════════════════════════════════════════════
# SECTION 3 — TECHNICAL FINDINGS
# ═══════════════════════════════════════════════════════════════

def _build_findings(doc, final_report: str, domains: list):
    _h1(doc, "3.", "Technical Findings")

    SKIP_RE = re.compile(
        r'executive summary|abstract|\u00f6zet|y\u00f6netici|'
        r'conclusion|sonu\u00e7|recommendation|\u00f6neri|'
        r'next step|reference|kaynaklar?|appendix|ek[\s:]|methodology',
        re.IGNORECASE
    )

    sections = _parse_sections(final_report)
    findings_sections = [
        (t, b) for (t, b) in sections
        if t and not SKIP_RE.search(t)
    ]

    if not findings_sections:
        _render_body(doc, final_report)
        return

    sub_counter = 1
    for title, body in findings_sections:
        _h2(doc, f"3.{sub_counter}", title)
        sub_counter += 1
        _render_body(doc, body)
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(4)


# ═══════════════════════════════════════════════════════════════
# SECTION 4 — DISCUSSION
# ═══════════════════════════════════════════════════════════════

def _build_discussion(doc, final_report: str, round_scores: list):
    """Section 4: Discussion — synthesised commentary and quality assessment."""
    _h1(doc, "4.", "Discussion")

    _h2(doc, "4.1", "Cross-Domain Synthesis")
    _body(doc, (
        "The domain analyses presented in Section 3 were subjected to numerical "
        "cross-validation to verify dimensional consistency, order-of-magnitude "
        "plausibility, and inter-agent coherence. Material properties, load "
        "parameters, and derived safety factors were checked across all active "
        "domains. Where agents raised cross-domain flags \u2014 indicating a finding "
        "in one discipline with consequences for another \u2014 these are annotated "
        "within the relevant subsection and must be resolved jointly before "
        "design can advance."
    ))

    rpn_img = chart_rpn(final_report)
    if rpn_img:
        _h2(doc, "4.2", "Risk Priority Summary")
        _body(doc, (
            "Failure modes identified across the domain analyses were ranked "
            "by Risk Priority Number (RPN = Severity \u00d7 Occurrence \u00d7 "
            "Detectability). The chart below provides a comparative view. "
            "Items with RPN \u2265 200 are classified Critical and require "
            "immediate design action; items in the range 100\u2013199 are High "
            "priority; items below 100 are Medium priority."
        ))
        _embed_chart(doc, rpn_img,
                     "Figure 4.1.  FMEA Risk Priority Numbers. "
                     "Critical: RPN \u2265 200  |  High: 100\u2013199  |  Medium: < 100.")

    if round_scores:
        _h2(doc, "4.3", "Quality Assessment")
        final_score = round_scores[-1].get("puan") if round_scores else None
        n_rounds    = len(round_scores)
        if final_score is not None:
            score_label = (
                "satisfactory (\u2265 85/100)" if final_score >= 85 else
                "acceptable (70\u201384/100)"  if final_score >= 70 else
                "below target (< 70/100)"
            )
            rounds_str = (
                f"across {n_rounds} analysis rounds"
                if n_rounds > 1 else "in a single analysis round"
            )
            if len(round_scores) > 1:
                progression = ", ".join(
                    f"Round {r['tur']}: {r.get('puan','\u2014')}/100"
                    for r in round_scores
                )
                prog_text = (
                    f" Score progression: {progression}. "
                    f"The {'improvement' if round_scores[-1].get('puan',0) > round_scores[0].get('puan',0) else 'consistency'} "
                    f"across rounds reflects the iterative feedback mechanism of the multi-agent pipeline."
                )
            else:
                prog_text = ""

            _body(doc, (
                f"The independent Observer meta-agent evaluated all domain outputs "
                f"{rounds_str} using a five-criterion weighted rubric: "
                f"technical accuracy (30%), internal consistency (25%), "
                f"assumption transparency (20%), analysis depth (15%), and "
                f"cross-validation quality (10%). "
                f"The final quality score is {final_score}/100 \u2014 {score_label}."
                f"{prog_text}"
            ))

            if len(round_scores) > 1:
                rows = []
                for r in round_scores:
                    p = r.get("puan", 0)
                    status = (
                        ("\u2713 Target met",    C_OK,   True)  if p >= 85 else
                        ("~  Acceptable",   C_WARN, False) if p >= 70 else
                        ("\u2717 Below target",  C_ERR,  True)
                    )
                    rows.append([f"Round {r['tur']}", f"{p} / 100", status])
                _make_table(doc, ["Round", "Quality Score", "Status"],
                            rows, [2.5, 3.5, 8.0],
                            caption="Table 4.1.  Observer quality scores by analysis round.")


# ═══════════════════════════════════════════════════════════════
# SECTION 5 — CONCLUSIONS & RECOMMENDATIONS
# ═══════════════════════════════════════════════════════════════

def _build_conclusions(doc, final_report: str):
    _h1(doc, "5.", "Conclusions and Recommendations")

    conc_text = _extract_section(
        final_report,
        [r"CONCLUSION", r"SONU\u00c7", r"KEY FINDINGS?"],
        stop_markers=["RECOMMENDATION", "NEXT STEP", "\u00d6NERI", "REFERENCE",
                      "KAYNAKLAR", "SUMMARY", "RISK ASSESSMENT"]
    )
    if conc_text:
        clean_lines = []
        for ln in conc_text.splitlines():
            stripped = ln.strip()
            if (re.search(r'[=\u2248]\s*[\d.,]+\s*[a-zA-Z/\u00b7]', stripped)
                    and stripped.count("=") >= 1
                    and len(stripped) < 200):
                continue
            clean_lines.append(ln)
        conc_text = "\n".join(clean_lines).strip()

    _h2(doc, "5.1", "Conclusions")
    if conc_text:
        _render_body(doc, conc_text)
    else:
        finding_lines = re.findall(
            r'(?:^|\n)(?:\d+\.\s+|[-\u2022]\s*)(.{40,180})(?=\n|$)',
            _strip_markup(final_report)
        )
        finding_lines = [
            l.strip() for l in finding_lines
            if not re.match(r'^[A-Z\s]{4,}$', l.strip())
            and len(l.strip().split()) > 8
            and not l.strip().startswith(("CRITICAL", "HIGH", "MEDIUM", "LOW"))
        ][:5]
        finding_lines = [
            l for l in finding_lines
            if not (re.search(r'[=\u2248]\s*[\d.,]+\s*[a-zA-Z]', l) and l.count("=") >= 1)
            and not re.match(r'^[A-Z]\s*[=:]\s*', l)
            and len(l.split()) >= 8
        ][:5]

        if finding_lines:
            _body(doc, (
                "The following principal conclusions are drawn from the "
                "domain analyses presented in Section 3:"
            ))
            for line in finding_lines:
                p = doc.add_paragraph()
                p.paragraph_format.space_after  = Pt(3)
                p.paragraph_format.left_indent  = Cm(0.7)
                p.paragraph_format.first_line_indent = Cm(-0.45)
                _run(p, "\u2022  ", bold=True, size=SZ_BODY, color=C_NAVY)
                _run(p, line.lstrip("-\u2022*0123456789.) "), size=SZ_BODY, color=C_BLACK)
                _set_line_spacing(p, 1.1)
        else:
            _body(doc, (
                "The principal technical conclusions of this analysis are "
                "presented in the domain-specific subsections of Section 3. "
                "Quantitative results, safety factors, and failure assessments "
                "constitute the analytical basis for the recommendations below."
            ))

    rec_text = _extract_section(
        final_report,
        [r"RECOMMENDATION", r"NEXT STEP", r"\u00d6NER\u0130"],
        stop_markers=["REFERENCE", "KAYNAKLAR", "APPENDIX"]
    )
    _h2(doc, "5.2", "Recommendations")
    if rec_text:
        _render_body(doc, rec_text)
    else:
        prio_lines = re.findall(
            r'^(CRITICAL|HIGH|MEDIUM|LOW):\s*(.+)$',
            final_report, re.MULTILINE | re.IGNORECASE
        )
        if prio_lines:
            for level, body_text in prio_lines:
                c_map = {"CRITICAL": C_ERR, "HIGH": C_WARN, "MEDIUM": C_NAVY, "LOW": C_MUTED}
                c = c_map.get(level.upper(), C_NAVY)
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Cm(0.4)
                _run(p, f"[{level.upper()}]  ", bold=True, size=SZ_BODY - 0.5, color=c)
                _run(p, body_text.strip(), size=SZ_BODY, color=C_BLACK)
                _set_line_spacing(p, 1.1)
        else:
            _body(doc, (
                "Specific recommendations are identified within the Technical "
                "Findings section. All recommendations are classified by priority: "
                "CRITICAL \u2014 must be addressed before design can advance; "
                "HIGH \u2014 should be resolved in the next design iteration; "
                "MEDIUM \u2014 address during the detailed design phase."
            ))


# ═══════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════

def _build_references(doc, final_report: str):
    _h1(doc, "", "References")

    ref_text = _extract_section(
        final_report,
        [r"REFERENCES?", r"KAYNAKLAR?"],
        stop_markers=["APPENDIX"]
    )
    if ref_text:
        lines = [l.strip() for l in ref_text.splitlines() if l.strip()]
        for j, line in enumerate(lines, 1):
            if not re.match(r'^\[?\d+\]?', line):
                line = f"[{j}]  {line}"
            p = doc.add_paragraph()
            p.paragraph_format.space_after        = Pt(4)
            p.paragraph_format.left_indent        = Cm(0.9)
            p.paragraph_format.first_line_indent  = Cm(-0.6)
            _set_line_spacing(p, 1.1)
            _run(p, line, size=SZ_SMALL, color=C_BLACK)
    else:
        _body(doc, (
            "References cited by the analysis agents are documented "
            "in-text within the Technical Findings section. Standard sources "
            "include ASM Handbooks, ASME/ASTM/MIL-SPEC standards, Eurocode "
            "series, and peer-reviewed engineering literature."
        ), spacing=1.0)


# ═══════════════════════════════════════════════════════════════
# APPENDIX A — ANALYSIS METADATA
# ═══════════════════════════════════════════════════════════════

def _build_appendix_a(doc, domains: list, mode: int, round_scores: list,
                      total_cost: float, kur: float, agent_log: list,
                      max_rounds: int):
    _appendix_h1(doc, "A", "Analysis Metadata")

    _h2(doc, "A.1", "Session Summary")

    mode_labels = {1: "Single Agent", 2: "Dual Agent",
                   3: "Semi-Automatic", 4: "Full Automatic"}
    n_domain_agents  = sum(1 for a in (agent_log or [])
                           if a.get("key","").endswith(("_a","_b")))
    n_support_agents = len(agent_log or []) - n_domain_agents

    rows = [
        ["Analysis Mode",        f"{mode}  \u2014  {mode_labels.get(mode,'\u2014')}"],
        ["Active Domains",       str(len(domains))],
        ["Rounds Completed",     str(len(round_scores)) if round_scores else "1"],
        ["Domain Agents",        str(n_domain_agents)],
        ["Support Agents",       str(n_support_agents)],
        ["Total Agents",         str(len(agent_log) if agent_log else 0)],
        ["Total Input Tokens",   "\u2014"],
        ["Total Output Tokens",  "\u2014"],
        ["Total API Cost (USD)", f"${total_cost:.5f}"],
        ["Approximate Cost (TL)", f"\u2248 {total_cost * kur:.2f} TL  (exchange rate: {kur:.1f})"],
    ]
    _make_table(doc, ["Parameter", "Value"], rows, [5.0, 10.0],
                caption="Table A.1.  Analysis session parameters and resource utilization.")

    if round_scores:
        _h2(doc, "A.2", "Quality Score History")
        img = chart_quality_scores(round_scores)
        if img:
            _embed_chart(doc, img,
                         "Figure A.1.  Observer quality score per analysis round. "
                         "Target threshold: 85/100. Score below 70/100 triggers mandatory revision.")
        else:
            score_rows = [
                [f"Round {r['tur']}", str(r.get('puan','\u2014')) + " / 100",
                 ("\u2713 Target met", C_OK, True)  if (r.get('puan') or 0) >= 85 else
                 ("~  Acceptable",C_WARN,False) if (r.get('puan') or 0) >= 70 else
                 ("\u2717 Below target",C_ERR,True)]
                for r in round_scores
            ]
            _make_table(doc, ["Round", "Quality Score", "Status"],
                        score_rows, [2.5, 3.5, 8.0],
                        caption="Table A.2.  Observer quality scores by round.")

    if agent_log:
        _h2(doc, "A.3", "Agent Cost Distribution")
        img = chart_agent_cost(agent_log)
        if img:
            _embed_chart(doc, img,
                         f"Figure A.2.  API cost per agent. Total: ${total_cost:.4f} USD "
                         f"(\u2248 {total_cost*kur:.2f} TL).")
        else:
            cost_rows = sorted(
                [[a.get("name","?"), f"${a.get('cost',0):.5f}"]
                 for a in agent_log if a.get("cost",0) > 0],
                key=lambda x: float(x[1][1:]), reverse=True
            )[:15]
            _make_table(doc, ["Agent", "Cost (USD)"], cost_rows, [11.0, 4.0],
                        caption="Table A.3.  Per-agent API cost.")


# ═══════════════════════════════════════════════════════════════
# APPENDIX B — AGENT ACTIVITY SUMMARY
# ═══════════════════════════════════════════════════════════════

def _summarise_output(text: str, max_words: int = 150) -> str:
    """Condense agent output to key sentences."""
    if not text or text.strip().startswith("ERROR"):
        return "(No output recorded.)"
    lines = [l.strip() for l in text.splitlines() if l.strip()
             and not re.match(r'^#{1,4}\s|^[=\-]{5,}$', l.strip())]
    joined = " ".join(lines)
    words  = joined.split()
    if len(words) <= max_words:
        return joined
    return " ".join(words[:max_words]) + " [...]"


def _build_appendix_b(doc, agent_log: list):
    if not agent_log:
        return

    _appendix_h1(doc, "B", "Agent Activity Summary")

    _body(doc, (
        "This appendix documents the outputs of all agents that participated in "
        "the analysis. Domain agent outputs are presented in full (condensed to "
        "preserve key findings); observer, cross-validation, and risk agent "
        "outputs are presented in full to allow independent review of the "
        "quality control process."
    ))

    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    try:
        import sys as _sys, os as _os
        _sys.path.insert(0, _os.path.join(_os.path.dirname(_os.path.dirname(__file__))))
        from config.agents_config import AGENTS as _DA
    except Exception:
        _DA = {}

    FULL_TEXT_KEYS = {
        "gozlemci", "capraz_dogrulama", "varsayim_belirsizlik",
        "risk_guvenilirlik", "celisiki_cozum", "soru_uretici",
        "literatur_patent"
    }

    def _model_tag(model_str):
        if "opus"   in (model_str or ""): return "Opus"
        if "sonnet" in (model_str or ""): return "Sonnet"
        if "haiku"  in (model_str or ""): return "Haiku"
        return model_str or "\u2014"

    domain_log   = [a for a in agent_log if a.get("key","") in _DA]
    qa_log       = [a for a in agent_log
                    if a.get("key","") in FULL_TEXT_KEYS]
    other_log    = [a for a in agent_log
                    if a.get("key","") not in _DA
                    and a.get("key","") not in FULL_TEXT_KEYS]

    if domain_log:
        _h2(doc, "B.1", "Domain Agent Outputs")
        _body(doc, (
            "Each domain agent conducted independent analysis from a specialist "
            "perspective. Outputs are condensed to key findings and numerical results."
        ))
        d_rows = [
            [a.get("name","?"),
             _model_tag(a.get("model","")),
             f"${a.get('cost',0):.5f}",
             _summarise_output(a.get("output",""), 250)]
            for a in domain_log
        ]
        _make_table(doc,
                    ["Agent", "Model", "Cost (USD)", "Key Findings (condensed)"],
                    d_rows, [3.5, 1.5, 2.0, 9.0],
                    caption="Table B.1.  Domain agent outputs \u2014 condensed to principal findings.")

    if qa_log:
        _h2(doc, "B.2", "Quality Assurance Agent Outputs")
        _body(doc, (
            "The following agents performed quality control, cross-validation, "
            "risk assessment, and critical questioning. Their outputs are "
            "presented in greater detail to allow independent review."
        ))
        for j, a in enumerate(qa_log, 1):
            name    = a.get("name","?")
            model   = _model_tag(a.get("model",""))
            cost    = a.get("cost",0)
            output  = (a.get("output") or "").strip()
            clean   = _strip_markup(output)
            words   = clean.split()
            if len(words) > 400:
                clean = " ".join(words[:400]) + " [condensed]"

            _h3(doc, f"B.2.{j}  {name}  \u00b7  {model}  \u00b7  ${cost:.5f}")
            if clean:
                _render_body(doc, clean, base_size=SZ_SMALL)
            sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    if other_log:
        _h2(doc, "B.3", "Other Support Agent Outputs")
        o_rows = [
            [a.get("name","?"),
             _model_tag(a.get("model","")),
             f"${a.get('cost',0):.5f}",
             _summarise_output(a.get("output",""), 80)]
            for a in other_log
        ]
        _make_table(doc,
                    ["Agent", "Model", "Cost (USD)", "Summary"],
                    o_rows, [4.0, 1.5, 2.0, 8.5],
                    caption="Table B.3.  Other support and synthesis agent outputs.")


# ═══════════════════════════════════════════════════════════════
# HEADER & FOOTER
# ═══════════════════════════════════════════════════════════════

def _setup_running_headers(doc, brief_short: str, doc_no: str):
    date_str = datetime.datetime.now().strftime("%Y-%m-%d")
    for section in doc.sections:
        hp = section.header.paragraphs[0]
        hp.clear()
        hp.paragraph_format.space_after = Pt(2)
        _run(hp, "Engineering AI  \u00b7  Technical Analysis Report", size=8, color=C_MUTED)
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

        fp = section.footer.paragraphs[0]
        fp.clear()
        _run(fp, f"{brief_short[:60]}{'...' if len(brief_short)>60 else ''}",
             size=7.5, color=C_MUTED, italic=True)
        _run(fp, f"    {date_str}    {doc_no}", size=7.5, color=C_MUTED)
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
