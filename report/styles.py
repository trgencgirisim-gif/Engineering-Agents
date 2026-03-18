"""report/styles.py — Typography constants, color constants, and low-level XML/DOCX helpers."""

import re
from typing import List

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ═══════════════════════════════════════════════════════════════
# TYPOGRAPHY & COLOR CONSTANTS — Academic / IEEE style
# ═══════════════════════════════════════════════════════════════

F          = "Times New Roman"   # body font
F_MONO     = "Courier New"       # code / numerical data

# Font sizes
SZ_BODY    = 12.0   # body text
SZ_H1      = 14.0   # numbered section heading
SZ_H2      = 12.0   # subsection
SZ_H3      = 11.0   # sub-subsection
SZ_CAPTION = 10.0   # figure/table captions
SZ_SMALL   = 10.0   # footnotes, metadata
SZ_ABSTRACT= 11.0   # abstract text

# Colors
C_BLACK    = RGBColor(0x00, 0x00, 0x00)
C_NAVY     = RGBColor(0x1F, 0x37, 0x62)   # IEEE dark blue
C_NAVY_HX  = "1F3762"
C_GREY     = RGBColor(0x44, 0x44, 0x44)
C_MUTED    = RGBColor(0x77, 0x77, 0x77)
C_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
C_OK       = RGBColor(0x1A, 0x7A, 0x4A)
C_WARN     = RGBColor(0x8A, 0x60, 0x00)
C_ERR      = RGBColor(0xC0, 0x44, 0x1E)

# Table
TH_BG      = C_NAVY_HX          # table header fill
ROW_ALT    = "F0F3F8"            # alternating row
BORDER_HX  = "B8C4D8"           # table border color
NOTE_BG    = "EEF3FA"            # note/info box fill
NOTE_BD    = "B0C4DE"            # note box border


# ═══════════════════════════════════════════════════════════════
# LOW-LEVEL XML HELPERS
# ═══════════════════════════════════════════════════════════════

def _shade(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def _borders(cell, color: str = BORDER_HX, sz: str = "4"):
    tcPr      = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _no_border_table(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        tblBorders.append(b)
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(tblBorders)


def _row_height(row, twips: int):
    trPr = row._tr.get_or_add_trPr()
    trH  = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),   str(twips))
    trH.set(qn("w:hRule"), "exact")
    trPr.append(trH)


def _set_line_spacing(para, spacing: float = 1.15):
    """Set line spacing (1.0, 1.15, 1.5, 2.0)."""
    pPr = para._p.get_or_add_pPr()
    spacing_el = OxmlElement("w:spacing")
    spacing_el.set(qn("w:line"),     str(int(spacing * 240)))
    spacing_el.set(qn("w:lineRule"), "auto")
    old = pPr.find(qn("w:spacing"))
    if old is not None:
        pPr.remove(old)
    pPr.append(spacing_el)


def _run(para, text, bold=False, italic=False, size=None,
         color=None, font=None, underline=False):
    run = para.add_run(text)
    run.bold    = bold
    run.italic  = italic
    run.underline = underline
    run.font.name = font or F
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def _set_margins(doc, top=2.5, right=2.0, bottom=2.5, left=2.5):
    """Academic margins: generous top/bottom, standard sides. Left wider for binding."""
    for section in doc.sections:
        section.page_width    = Cm(21.0)
        section.page_height   = Cm(29.7)
        section.top_margin    = Cm(top)
        section.right_margin  = Cm(right)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)


# ═══════════════════════════════════════════════════════════════
# DOCUMENT-LEVEL STYLE SETUP
# ═══════════════════════════════════════════════════════════════

def _setup_styles(doc):
    """Apply Times New Roman as document default. Override built-in heading styles."""
    try:
        normal = doc.styles["Normal"]
        normal.font.name = F
        normal.font.size = Pt(SZ_BODY)
        normal.paragraph_format.space_after = Pt(6)
    except Exception:
        pass

    heading_defs = [
        ("Heading 1", SZ_H1,  True,  False, 24, 8),
        ("Heading 2", SZ_H2,  True,  False, 16, 4),
        ("Heading 3", SZ_H3,  True,  True,  12, 3),
    ]
    for name, sz, bold, italic, sb, sa in heading_defs:
        try:
            s = doc.styles[name]
            s.font.name  = F
            s.font.size  = Pt(sz)
            s.font.bold  = bold
            s.font.italic = italic
            s.font.color.rgb = C_BLACK
            s.paragraph_format.space_before = Pt(sb)
            s.paragraph_format.space_after  = Pt(sa)
            s.paragraph_format.keep_with_next = True
        except Exception:
            pass


# ═══════════════════════════════════════════════════════════════
# BUILDING BLOCKS
# ═══════════════════════════════════════════════════════════════

def _h_rule(doc, color: str = "CCCCCC", thick: int = 8):
    """Horizontal rule as a thin shaded table row."""
    t = doc.add_table(rows=1, cols=1)
    _no_border_table(t)
    _shade(t.rows[0].cells[0], color)
    t.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Pt(0)
    t.rows[0].cells[0].paragraphs[0].paragraph_format.space_after  = Pt(0)
    _row_height(t.rows[0], thick)


def _light_rule(doc):
    _h_rule(doc, "DDDDDD", 6)


def _h1(doc, number: str, title: str):
    """Numbered section heading: '1. INTRODUCTION'"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    _run(p, f"{number}  ", bold=True, size=SZ_H1, color=C_NAVY)
    _run(p, title.upper(), bold=True, size=SZ_H1, color=C_BLACK)
    _h_rule(doc, color=C_NAVY_HX, thick=12)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


def _h2(doc, number: str, title: str):
    """Subsection heading: '1.1  Materials Analysis'"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.keep_with_next = True
    _run(p, f"{number}  ", bold=True, size=SZ_H2, color=C_NAVY)
    _run(p, title, bold=True, size=SZ_H2, color=C_BLACK)


def _h3(doc, title: str):
    """Sub-subsection heading — italic, no number."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    _run(p, title, bold=True, italic=True, size=SZ_H3, color=C_GREY)


def _appendix_h1(doc, letter: str, title: str):
    """Appendix section heading: 'APPENDIX A — ANALYSIS METADATA'"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)
    _run(p, f"APPENDIX {letter}  —  ", bold=True, size=SZ_H1, color=C_NAVY)
    _run(p, title.upper(), bold=True, size=SZ_H1, color=C_BLACK)
    _h_rule(doc, color=C_NAVY_HX, thick=12)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


def _body(doc, text: str, spacing: float = 1.15, indent: float = 0):
    """Body paragraph with proper academic line spacing."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    _set_line_spacing(p, spacing)
    _run(p, text, size=SZ_BODY, color=C_BLACK)


def _caption(doc, text: str, above: bool = True):
    """Figure or table caption in small italic."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    _run(p, text, italic=True, size=SZ_CAPTION, color=C_MUTED)


def _note_box(doc, label: str, body: str):
    """Shaded note/flag box for cross-domain flags, assumptions, etc."""
    t = doc.add_table(rows=1, cols=1)
    _no_border_table(t)
    cell = t.rows[0].cells[0]
    _shade(cell, NOTE_BG)
    _borders(cell, NOTE_BD, "4")

    lp = cell.paragraphs[0]
    lp.paragraph_format.space_before = Pt(5)
    lp.paragraph_format.space_after  = Pt(2)
    lp.paragraph_format.left_indent  = Cm(0.3)
    _run(lp, label + "  ", bold=True, size=SZ_SMALL, color=C_NAVY)
    _run(lp, body,         bold=False, size=SZ_SMALL, color=C_GREY)

    bot = cell.add_paragraph()
    bot.paragraph_format.space_after = Pt(4)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


def _make_table(doc, headers: list, rows: list, widths_cm: list,
                header_bg: str = TH_BG, caption: str = ""):
    """Academic-style table: header row + alternating body rows."""
    if caption:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(8)
        cp.paragraph_format.space_after  = Pt(2)
        _run(cp, caption, italic=True, size=SZ_CAPTION, color=C_MUTED)

    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"

    hr = tbl.rows[0]
    for i, hdr in enumerate(headers):
        cell = hr.cells[i]
        _shade(cell, header_bg)
        _borders(cell, BORDER_HX)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.left_indent  = Cm(0.2)
        _run(p, hdr, bold=True, size=SZ_SMALL, color=C_WHITE)

    for ri, row_data in enumerate(rows):
        dr = tbl.rows[ri + 1]
        bg = "FFFFFF" if ri % 2 == 0 else ROW_ALT
        for ci, val in enumerate(row_data):
            cell = dr.cells[ci]
            _shade(cell, bg)
            _borders(cell, BORDER_HX)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Cm(0.2)
            if isinstance(val, tuple):
                text  = val[0]
                color = val[1]
                bold  = val[2] if len(val) > 2 else False
                _run(p, str(text), bold=bool(bold), size=SZ_SMALL, color=color)
            else:
                _run(p, str(val), size=SZ_SMALL, color=C_BLACK)

    for row in tbl.rows:
        for ci, cell in enumerate(row.cells):
            if ci < len(widths_cm):
                cell.width = Cm(widths_cm[ci])

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)
    return tbl


# ═══════════════════════════════════════════════════════════════
# CONTENT PARSER & RENDERER
# ═══════════════════════════════════════════════════════════════

def _strip_markup(text: str) -> str:
    """Remove markdown headings, bold markers, excess whitespace from text."""
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'^[|\-\s:]{4,}$', '', text, flags=re.MULTILINE)
    text = re.sub(r'^\|.*\|$', '', text, flags=re.MULTILINE)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def _extract_abstract(text: str, max_sentences: int = 5) -> str:
    """Extract a clean abstract from final_report."""
    for marker in ["EXECUTIVE SUMMARY", "ABSTRACT", "ÖZET", "YÖNETİCİ ÖZETİ"]:
        m = re.search(
            rf'(?:^|\n)(?:#{1,4}\s*)?{marker}[^\n]*\n+(.*?)(?=\n(?:#{1,4}\s+|\d+\.\s+[A-Z]|[A-Z]{{4,}}\n)|\Z)',
            text, re.IGNORECASE | re.DOTALL
        )
        if m:
            raw = m.group(1)
            raw = re.split(r'\n#{1,4}\s+|\n\d+\.\s+[A-Z]', raw)[0]
            content = _strip_markup(raw).strip()
            content = _to_prose_sentences(content, max_sentences)
            if len(content.split()) > 25:
                return content

    clean = _strip_markup(text)
    paras = []
    for para in re.split(r'\n{2,}', clean):
        para = para.strip()
        if (len(para.split()) < 25
                or re.match(r'^[A-Z\s&.:0-9/-]{6,}$', para)
                or re.match(r'^[-•*\d]', para)):
            continue
        paras.append(para)
    if paras:
        combined = " ".join(paras[:2])
        return _to_prose_sentences(combined, max_sentences)
    return ""


def _to_prose_sentences(text: str, max_sentences: int = 5) -> str:
    """Clean text to up to max_sentences complete sentences."""
    lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            lines.append("")
            continue
        if (re.search(r'[=≈]\s*[\d.,]+\s*[a-zA-Z]', stripped)
                and len(stripped) < 150
                and stripped.count("=") >= 1):
            continue
        if re.match(r'^[-•*–·]\s|^\d+[.)]', stripped):
            continue
        if len(stripped.split()) < 5:
            continue
        lines.append(stripped)

    prose = " ".join(l for l in lines if l).strip()
    prose = re.sub(r'\s{2,}', ' ', prose)

    # Protect common abbreviations from sentence splitting
    _ABBREVS = ["e.g.", "i.e.", "vs.", "etc.", "Dr.", "Mr.", "Ms.",
                "Prof.", "Fig.", "Eq.", "No.", "Vol.", "approx.",
                "min.", "max.", "avg.", "ref.", "incl."]
    for a in _ABBREVS:
        prose = prose.replace(a, a.replace(".", "\x00"))

    sentences = re.split(r'(?<=[.!?])\s+', prose)
    sentences = [s.replace("\x00", ".") for s in sentences if len(s.split()) >= 3]

    result = " ".join(sentences[:max_sentences])
    if result and result[-1] not in ".!?":
        result += "."
    return result


def _parse_sections(text: str) -> List[tuple]:
    """Parse final_report text into (title, body) tuples."""
    lines_raw = text.splitlines()
    joined = []
    j = 0
    while j < len(lines_raw):
        line = lines_raw[j]
        stripped = line.strip()
        if "CROSS-DOMAIN FLAG" in stripped.upper() and not stripped.endswith("."):
            combined = stripped
            j += 1
            while j < len(lines_raw):
                nxt = lines_raw[j].strip()
                if (not nxt
                        or re.match(r'^#{1,4}\s', nxt)
                        or re.match(r'^\d+\.\s+[A-Z]', nxt)
                        or re.match(r'^[A-Z][A-Z &/:0-9-]{4,}$', nxt)):
                    break
                combined = combined + " " + nxt
                j += 1
            joined.append(combined)
        else:
            joined.append(line)
            j += 1
    text = "\n".join(joined)

    heading_re = re.compile(
        r'^(?:'
        r'#{1,4}\s+.{3,}'
        r'|(?:\d+\.?\d*\.?\s+[A-Z].{2,})'
        r'|(?:\*\*[A-Z].{3,}\*\*)'
        r'|(?:[A-Z][A-Z &/:0-9\-]{4,})$'
        r')'
    )
    sections, cur_title, cur_lines = [], "", []
    for line in text.splitlines():
        s = line.strip()
        if s and heading_re.match(s):
            body = "\n".join(cur_lines).strip()
            if body:
                sections.append((cur_title, body))
            cur_title = re.sub(r'^[#\s*0-9.]+', '', s).strip().strip("*:").strip()
            cur_lines = []
        else:
            cur_lines.append(line)
    body = "\n".join(cur_lines).strip()
    if body:
        sections.append((cur_title, body))
    return sections


def _clean_heading(text: str) -> str:
    return re.sub(r'^#{1,6}\s*', '', text).strip().strip("*:").strip()


def _is_md_heading(text: str) -> bool:
    return bool(re.match(r'^#{1,6}\s+.{2,}', text))


def _is_section_divider(text: str) -> bool:
    return bool(re.match(r'^[=\-]{4,}\s*$', text))


def _render_body(doc, text: str, base_size: float = SZ_BODY):
    """Render free-form agent output into the document with academic typography."""
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i].strip()
        i += 1

        if not raw:
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(3)
            continue

        if _is_section_divider(raw):
            continue

        if _is_md_heading(raw):
            heading_text = _clean_heading(raw)
            if not heading_text:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.keep_with_next = True
            _run(p, heading_text, bold=True, italic=True,
                 size=base_size, color=C_NAVY)
            continue

        if (re.match(r'^[A-Z][A-Z &/:0-9-]{4,}$', raw)
                and not re.search(r'[=<>]', raw)
                and len(raw) < 60):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.keep_with_next = True
            _run(p, raw.title(), bold=True, italic=True,
                 size=base_size, color=C_NAVY)
            continue

        if raw.startswith("|") and raw.count("|") >= 2:
            if re.match(r'^[|\-\s:]+$', raw):
                continue
            cells = [c.strip() for c in raw.strip("|").split("|")]
            if all(len(c) <= 2 for c in cells):
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.4)
            _run(p, "   |   ".join(cells), size=base_size - 1.5,
                 color=C_GREY, font=F_MONO)
            continue

        if re.match(r'^[-•*–·]\s', raw):
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Cm(0.7)
            p.paragraph_format.first_line_indent = Cm(-0.45)
            _run(p, "•  ", bold=True, size=base_size, color=C_NAVY)
            _run(p, raw.lstrip("-•*–· ").strip(), size=base_size, color=C_BLACK)
            _set_line_spacing(p, 1.1)
            continue

        m_num = re.match(r'^(\d+[.)]\s+)(.*)', raw)
        if m_num:
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Cm(0.7)
            p.paragraph_format.first_line_indent = Cm(-0.45)
            _run(p, m_num.group(1).rstrip() + "  ", bold=True,
                 size=base_size, color=C_NAVY)
            _run(p, m_num.group(2).strip(), size=base_size, color=C_BLACK)
            _set_line_spacing(p, 1.1)
            continue

        m_prio = re.match(r'^(CRITICAL|HIGH|MEDIUM|LOW):\s*(.*)', raw, re.IGNORECASE)
        if m_prio:
            level     = m_prio.group(1).upper()
            body_text = m_prio.group(2).strip()
            c_map = {"CRITICAL": C_ERR, "HIGH": C_WARN, "MEDIUM": C_NAVY, "LOW": C_MUTED}
            c = c_map.get(level, C_NAVY)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.4)
            _run(p, f"[{level}]  ", bold=True, size=base_size - 0.5, color=c)
            _run(p, body_text, size=base_size, color=C_BLACK)
            _set_line_spacing(p, 1.1)
            continue

        if "CROSS-DOMAIN FLAG" in raw.upper():
            body_clean = re.sub(r'CROSS[-\s]?DOMAIN\s+FLAG[:\s]*',
                                '', raw, flags=re.IGNORECASE).strip()
            _note_box(doc, "\u2691  Cross-Domain Flag", body_clean)
            continue

        if (re.search(r'[=\u2248]\s*[\d.,]+\s*[a-zA-Z]', raw)
                and not raw.endswith(":")
                and len(raw) < 220):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.left_indent  = Cm(1.4)
            _run(p, raw, size=base_size - 1, color=C_BLACK, font=F_MONO)
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _set_line_spacing(p, 1.15)
        parts = re.split(r'(\*\*[^*]+\*\*)', raw)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                _run(p, part[2:-2], bold=True, size=base_size, color=C_BLACK)
            else:
                _run(p, part, size=base_size, color=C_BLACK)


def _extract_section(text: str, markers: list, stop_markers: list = None) -> str:
    """Extract a named section from report text."""
    stop_pat = stop_markers or []
    stop_group = (
        r'(?=\n#{1,4}\s|\n\d+\.\s+[A-Z]|\n[A-Z]{4,}\n'
        + (r'|\n' + r'|\n'.join(re.escape(s) for s in stop_pat) if stop_pat else r'')
        + r'|\Z)'
    )
    for marker in markers:
        pattern = rf'(?:^|\n){marker}[^\n]*\n+(.*?)' + stop_group
        m = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if m and len(m.group(1).strip()) > 30:
            return m.group(1).strip()
    return ""
