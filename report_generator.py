"""
Engineering AI — Report Generator v2
Academic-style DOCX report. Times New Roman, IEEE/ASME conventions.
Structure: Cover → Abstract → Introduction → Methodology → Findings
           → Discussion → Conclusions → References → Appendix A (Metadata) → Appendix B (Agent Summaries)
"""

import io
import re
import datetime
from typing import List, Dict, Optional
from collections import defaultdict

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.shared import OxmlElement as OxmlEl

try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    import matplotlib.ticker as mticker
    MATPLOTLIB_OK = True
except ImportError:
    MATPLOTLIB_OK = False

# ═══════════════════════════════════════════════════════════════
# TYPOGRAPHY & COLOR CONSTANTS — Academic / IEEE style
# ═══════════════════════════════════════════════════════════════

F          = "Times New Roman"   # body font
F_MONO     = "Courier New"       # code / numerical data

# Font sizes
SZ_BODY    = 12.0   # body text
SZ_H1      = 14.0   # numbered section heading
SZ_H2      = 12.0   # subsection
SZ_H3      = 11.0   # sub-subsection
SZ_CAPTION = 10.0   # figure/table captions
SZ_SMALL   = 10.0   # footnotes, metadata
SZ_ABSTRACT= 11.0   # abstract text

# Colors
C_BLACK    = RGBColor(0x00, 0x00, 0x00)
C_NAVY     = RGBColor(0x1F, 0x37, 0x62)   # IEEE dark blue
C_NAVY_HX  = "1F3762"
C_GREY     = RGBColor(0x44, 0x44, 0x44)
C_MUTED    = RGBColor(0x77, 0x77, 0x77)
C_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
C_OK       = RGBColor(0x1A, 0x7A, 0x4A)
C_WARN     = RGBColor(0x8A, 0x60, 0x00)
C_ERR      = RGBColor(0xC0, 0x44, 0x1E)

# Table
TH_BG      = C_NAVY_HX          # table header fill
ROW_ALT    = "F0F3F8"            # alternating row
BORDER_HX  = "B8C4D8"           # table border color
NOTE_BG    = "EEF3FA"            # note/info box fill
NOTE_BD    = "B0C4DE"            # note box border


# ═══════════════════════════════════════════════════════════════
# LOW-LEVEL XML HELPERS
# ═══════════════════════════════════════════════════════════════

def _shade(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def _borders(cell, color: str = BORDER_HX, sz: str = "4"):
    tcPr      = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _no_border_table(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        tblBorders.append(b)
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(tblBorders)


def _row_height(row, twips: int):
    trPr = row._tr.get_or_add_trPr()
    trH  = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),   str(twips))
    trH.set(qn("w:hRule"), "exact")
    trPr.append(trH)


def _set_line_spacing(para, spacing: float = 1.15):
    """Set line spacing (1.0, 1.15, 1.5, 2.0)."""
    pPr = para._p.get_or_add_pPr()
    spacing_el = OxmlElement("w:spacing")
    spacing_el.set(qn("w:line"),     str(int(spacing * 240)))
    spacing_el.set(qn("w:lineRule"), "auto")
    old = pPr.find(qn("w:spacing"))
    if old is not None:
        pPr.remove(old)
    pPr.append(spacing_el)


def _run(para, text, bold=False, italic=False, size=None,
         color=None, font=None, underline=False):
    run = para.add_run(text)
    run.bold    = bold
    run.italic  = italic
    run.underline = underline
    run.font.name = font or F
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def _set_margins(doc, top=2.5, right=2.0, bottom=2.5, left=2.5):
    """Academic margins: generous top/bottom, standard sides. Left wider for binding."""
    for section in doc.sections:
        section.page_width    = Cm(21.0)
        section.page_height   = Cm(29.7)
        section.top_margin    = Cm(top)
        section.right_margin  = Cm(right)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)


# ═══════════════════════════════════════════════════════════════
# DOCUMENT-LEVEL STYLE SETUP
# ═══════════════════════════════════════════════════════════════

def _setup_styles(doc):
    """Apply Times New Roman as document default. Override built-in heading styles."""
    try:
        normal = doc.styles["Normal"]
        normal.font.name = F
        normal.font.size = Pt(SZ_BODY)
        normal.paragraph_format.space_after = Pt(6)
    except Exception:
        pass

    # Override built-in headings for TOC compatibility
    heading_defs = [
        ("Heading 1", SZ_H1,  True,  False, 24, 8),
        ("Heading 2", SZ_H2,  True,  False, 16, 4),
        ("Heading 3", SZ_H3,  True,  True,  12, 3),
    ]
    for name, sz, bold, italic, sb, sa in heading_defs:
        try:
            s = doc.styles[name]
            s.font.name  = F
            s.font.size  = Pt(sz)
            s.font.bold  = bold
            s.font.italic = italic
            s.font.color.rgb = C_BLACK
            s.paragraph_format.space_before = Pt(sb)
            s.paragraph_format.space_after  = Pt(sa)
            s.paragraph_format.keep_with_next = True
        except Exception:
            pass


# ═══════════════════════════════════════════════════════════════
# BUILDING BLOCKS
# ═══════════════════════════════════════════════════════════════

def _h1(doc, number: str, title: str):
    """Numbered section heading: '1. INTRODUCTION'"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    # Section number
    r1 = _run(p, f"{number}  ", bold=True, size=SZ_H1, color=C_NAVY)
    # Section title in small caps style (bold uppercase)
    r2 = _run(p, title.upper(), bold=True, size=SZ_H1, color=C_BLACK)
    # Rule below
    _h_rule(doc, color=C_NAVY_HX, thick=12)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


def _h2(doc, number: str, title: str):
    """Subsection heading: '1.1  Materials Analysis'"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.keep_with_next = True
    _run(p, f"{number}  ", bold=True, size=SZ_H2, color=C_NAVY)
    _run(p, title, bold=True, size=SZ_H2, color=C_BLACK)


def _h3(doc, title: str):
    """Sub-subsection heading — italic, no number (inline narrative headings)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    _run(p, title, bold=True, italic=True, size=SZ_H3, color=C_GREY)


def _appendix_h1(doc, letter: str, title: str):
    """Appendix section heading: 'APPENDIX A — ANALYSIS METADATA'"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)
    _run(p, f"APPENDIX {letter}  —  ", bold=True, size=SZ_H1, color=C_NAVY)
    _run(p, title.upper(), bold=True, size=SZ_H1, color=C_BLACK)
    _h_rule(doc, color=C_NAVY_HX, thick=12)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


def _h_rule(doc, color: str = "CCCCCC", thick: int = 8):
    """Horizontal rule as a thin shaded table row."""
    t = doc.add_table(rows=1, cols=1)
    _no_border_table(t)
    _shade(t.rows[0].cells[0], color)
    t.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Pt(0)
    t.rows[0].cells[0].paragraphs[0].paragraph_format.space_after  = Pt(0)
    _row_height(t.rows[0], thick)


def _light_rule(doc):
    _h_rule(doc, "DDDDDD", 6)


def _body(doc, text: str, spacing: float = 1.15, indent: float = 0):
    """Body paragraph with proper academic line spacing."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    _set_line_spacing(p, spacing)
    _run(p, text, size=SZ_BODY, color=C_BLACK)


def _caption(doc, text: str, above: bool = True):
    """Figure or table caption in small italic."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    _run(p, text, italic=True, size=SZ_CAPTION, color=C_MUTED)


def _note_box(doc, label: str, body: str):
    """Shaded note/flag box for cross-domain flags, assumptions, etc."""
    t = doc.add_table(rows=1, cols=1)
    _no_border_table(t)
    cell = t.rows[0].cells[0]
    _shade(cell, NOTE_BG)
    _borders(cell, NOTE_BD, "4")

    # Label
    lp = cell.paragraphs[0]
    lp.paragraph_format.space_before = Pt(5)
    lp.paragraph_format.space_after  = Pt(2)
    lp.paragraph_format.left_indent  = Cm(0.3)
    _run(lp, label + "  ", bold=True, size=SZ_SMALL, color=C_NAVY)
    _run(lp, body,         bold=False, size=SZ_SMALL, color=C_GREY)

    bot = cell.add_paragraph()
    bot.paragraph_format.space_after = Pt(4)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


def _make_table(doc, headers: list, rows: list, widths_cm: list,
                header_bg: str = TH_BG, caption: str = ""):
    """Academic-style table: header row + alternating body rows."""
    if caption:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(8)
        cp.paragraph_format.space_after  = Pt(2)
        _run(cp, caption, italic=True, size=SZ_CAPTION, color=C_MUTED)

    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"

    # Header
    hr = tbl.rows[0]
    for i, hdr in enumerate(headers):
        cell = hr.cells[i]
        _shade(cell, header_bg)
        _borders(cell, BORDER_HX)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.left_indent  = Cm(0.2)
        _run(p, hdr, bold=True, size=SZ_SMALL, color=C_WHITE)

    # Body
    for ri, row_data in enumerate(rows):
        dr = tbl.rows[ri + 1]
        bg = "FFFFFF" if ri % 2 == 0 else ROW_ALT
        for ci, val in enumerate(row_data):
            cell = dr.cells[ci]
            _shade(cell, bg)
            _borders(cell, BORDER_HX)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Cm(0.2)
            if isinstance(val, tuple):
                text  = val[0]
                color = val[1]
                bold  = val[2] if len(val) > 2 else False
                _run(p, str(text), bold=bool(bold), size=SZ_SMALL, color=color)
            else:
                _run(p, str(val), size=SZ_SMALL, color=C_BLACK)

    # Column widths
    for row in tbl.rows:
        for ci, cell in enumerate(row.cells):
            if ci < len(widths_cm):
                cell.width = Cm(widths_cm[ci])

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)
    return tbl


# ═══════════════════════════════════════════════════════════════
# CHARTS — kept from v1, font updated
# ═══════════════════════════════════════════════════════════════

def _fig_bytes(fig) -> bytes:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                facecolor="white", edgecolor="none")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _embed_chart(doc, img_bytes: bytes, caption: str = ""):
    if not img_bytes:
        return
    doc.add_picture(io.BytesIO(img_bytes), width=Inches(5.5))
    if caption:
        _caption(doc, caption)


def chart_quality_scores(round_scores: List[Dict]) -> Optional[bytes]:
    if not MATPLOTLIB_OK:
        return None
    puanlar = [r["puan"] for r in round_scores if r.get("puan") is not None]
    if not puanlar:
        return None
    xs = list(range(1, len(puanlar) + 1))
    fig, ax = plt.subplots(figsize=(5.5, 2.8))
    ax.set_facecolor("white"); fig.patch.set_facecolor("white")
    ax.fill_between(xs, puanlar, alpha=0.1, color="#1F3762")
    ax.plot(xs, puanlar, color="#1F3762", linewidth=2, zorder=3)
    ax.scatter(xs, puanlar, color="white", edgecolors="#1F3762", s=55, linewidths=2, zorder=4)
    ax.axhline(85, color="#1A7A4A", linestyle="--", linewidth=1, alpha=0.7, label="Target (85)")
    ax.axhline(70, color="#8A6000", linestyle=":",  linewidth=1, alpha=0.7, label="Minimum (70)")
    for x, v in zip(xs, puanlar):
        c = "#1A7A4A" if v >= 85 else "#8A6000" if v >= 70 else "#C0441E"
        ax.annotate(str(v), (x, v), textcoords="offset points",
                    xytext=(0, 8), ha="center", fontsize=8, color=c, fontweight="bold",
                    fontfamily="serif")
    ax.set_xticks(xs)
    ax.set_xticklabels([f"Round {i}" for i in xs], fontsize=8, fontfamily="serif")
    ax.set_yticks([0, 25, 50, 75, 100])
    ax.set_ylim(0, 115)
    ax.tick_params(labelsize=8)
    ax.grid(axis="y", linewidth=0.4, color="#EEEEEE")
    ax.spines[["top", "right"]].set_visible(False)
    ax.legend(fontsize=7.5, loc="lower right", framealpha=0.7,
              prop={"family": "serif"})
    ax.set_title("Figure A.1.  Observer Quality Score per Analysis Round",
                 fontsize=9, fontweight="bold", color="#1A1A1A", pad=8,
                 fontfamily="serif", loc="left")
    fig.tight_layout()
    return _fig_bytes(fig)


def chart_agent_cost(agent_log: List[Dict], top_n: int = 12) -> Optional[bytes]:
    if not MATPLOTLIB_OK:
        return None
    agg = defaultdict(float)
    for a in agent_log:
        cost = a.get("cost", 0)
        if cost > 0:
            agg[a.get("name", a.get("key", "?"))[:32]] += cost
    if not agg:
        return None
    items  = sorted(agg.items(), key=lambda x: x[1], reverse=True)[:top_n]
    names  = [x[0] for x in items]
    costs  = [x[1] for x in items]
    max_c  = max(costs)
    colors = ["#1F3762" if c == max_c else "#6A8BC0" for c in costs]
    fig, ax = plt.subplots(figsize=(5.5, max(2.5, len(names) * 0.35 + 0.8)))
    ax.set_facecolor("white"); fig.patch.set_facecolor("white")
    ax.barh(range(len(names)), costs, color=colors, edgecolor="none", height=0.6)
    for i, cost in enumerate(costs):
        ax.text(cost + max_c * 0.01, i, f"${cost:.5f}",
                va="center", fontsize=7, color="#444444", fontfamily="serif")
    ax.set_yticks(range(len(names)))
    ax.set_yticklabels(names, fontsize=7.5, fontfamily="serif")
    ax.invert_yaxis()
    ax.xaxis.set_major_formatter(mticker.FormatStrFormatter("$%.4f"))
    ax.tick_params(labelsize=7.5)
    ax.grid(axis="x", linewidth=0.4, color="#EEEEEE")
    ax.spines[["top", "right"]].set_visible(False)
    ax.set_title(f"Figure A.2.  API Cost Distribution by Agent (Top {len(names)})",
                 fontsize=9, fontweight="bold", color="#1A1A1A", pad=8,
                 fontfamily="serif", loc="left")
    fig.tight_layout()
    return _fig_bytes(fig)


def chart_rpn(report_text: str) -> Optional[bytes]:
    if not MATPLOTLIB_OK:
        return None
    rpn_vals = []
    for label, rpn_str in re.findall(
        r'([A-Za-z][^\n]{3,50}?)\s*(?:RPN|rpn)\s*[=:]\s*(\d{2,3})', report_text
    )[:10]:
        rpn_vals.append((label.strip()[:35], int(rpn_str)))
    if not rpn_vals:
        for m in re.findall(r'(\d{1,2})\s*[×xX*]\s*(\d{1,2})\s*[×xX*]\s*(\d{1,2})', report_text)[:8]:
            s, o, d = int(m[0]), int(m[1]), int(m[2])
            rpn_vals.append((f"FM: S{s}×O{o}×D{d}", s * o * d))
    if not rpn_vals:
        return None
    labels = [v[0] for v in rpn_vals]
    rpns   = [v[1] for v in rpn_vals]
    max_r  = max(rpns)
    colors = ["#C0441E" if r >= 200 else "#C47A00" if r >= 100 else "#1A7A4A" for r in rpns]
    fig, ax = plt.subplots(figsize=(5.5, max(2.5, len(labels) * 0.38 + 1.2)))
    ax.set_facecolor("white"); fig.patch.set_facecolor("white")
    ax.barh(range(len(labels)), rpns, color=colors, edgecolor="none", height=0.6)
    if max_r >= 200:
        ax.axvline(200, color="#C0441E", linestyle="--", linewidth=1, alpha=0.5, label="Critical ≥200")
    if max_r >= 100:
        ax.axvline(100, color="#C47A00", linestyle=":",  linewidth=1, alpha=0.5, label="High 100–199")
    for i, rpn in enumerate(rpns):
        ax.text(rpn + max_r * 0.01, i, str(rpn), va="center", fontsize=8,
                color="#333333", fontfamily="serif")
    ax.set_yticks(range(len(labels)))
    ax.set_yticklabels(labels, fontsize=7.5, fontfamily="serif")
    ax.invert_yaxis()
    ax.tick_params(labelsize=7.5)
    ax.grid(axis="x", linewidth=0.4, color="#EEEEEE")
    ax.spines[["top", "right"]].set_visible(False)
    legend = [
        mpatches.Patch(color="#C0441E", label="Critical (RPN ≥ 200)"),
        mpatches.Patch(color="#C47A00", label="High (100 ≤ RPN < 200)"),
        mpatches.Patch(color="#1A7A4A", label="Medium (RPN < 100)"),
    ]
    ax.legend(handles=legend, fontsize=7, loc="lower right", framealpha=0.7,
              prop={"family": "serif"})
    ax.set_title("Figure 4.1.  FMEA Risk Priority Numbers",
                 fontsize=9, fontweight="bold", color="#1A1A1A", pad=8,
                 fontfamily="serif", loc="left")
    fig.tight_layout()
    return _fig_bytes(fig)


# ═══════════════════════════════════════════════════════════════
# CONTENT PARSER & RENDERER
# ═══════════════════════════════════════════════════════════════

def _strip_markup(text: str) -> str:
    """Remove markdown headings, bold markers, excess whitespace from text."""
    # Remove ## headings (keep content)
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    # Remove **bold**
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    # Remove table separator rows
    text = re.sub(r'^[|\-\s:]{4,}$', '', text, flags=re.MULTILINE)
    # Remove bare | table rows (replace with spaces)
    text = re.sub(r'^\|.*\|$', '', text, flags=re.MULTILINE)
    # Collapse multiple blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def _extract_abstract(text: str, max_sentences: int = 5) -> str:
    """
    Extract a clean abstract from final_report.
    Returns up to max_sentences clean sentences.
    Strips all markdown, headings, bullets, equations from result.
    """
    # ── Step 1: Look for named section ──────────────────────
    for marker in ["EXECUTIVE SUMMARY", "ABSTRACT", "ÖZET", "YÖNETİCİ ÖZETİ"]:
        m = re.search(
            rf'(?:^|\n)(?:#{1,4}\s*)?{marker}[^\n]*\n+(.*?)(?=\n(?:#{1,4}\s+|\d+\.\s+[A-Z]|[A-Z]{{4,}}\n)|\Z)',
            text, re.IGNORECASE | re.DOTALL
        )
        if m:
            raw = m.group(1)
            # Stop at first heading-like line within the section
            raw = re.split(r'\n#{1,4}\s+|\n\d+\.\s+[A-Z]', raw)[0]
            content = _strip_markup(raw).strip()
            content = _to_prose_sentences(content, max_sentences)
            if len(content.split()) > 25:
                return content

    # ── Step 2: Fallback — first substantive prose paragraph ─
    clean = _strip_markup(text)
    # Split into paragraphs, keep only prose (no all-caps, no lists)
    paras = []
    for para in re.split(r'\n{2,}', clean):
        para = para.strip()
        if (len(para.split()) < 25
                or re.match(r'^[A-Z\s&.:0-9/-]{6,}$', para)
                or re.match(r'^[-•*\d]', para)):
            continue
        paras.append(para)
    if paras:
        combined = " ".join(paras[:2])
        return _to_prose_sentences(combined, max_sentences)
    return ""


def _to_prose_sentences(text: str, max_sentences: int = 5) -> str:
    """
    Clean text → up to max_sentences complete sentences.
    Removes list markers, equation lines, and dangling fragments.
    """
    # Remove list-marker lines
    lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            lines.append("")
            continue
        # Skip lines that are pure equations/calculations
        if (re.search(r'[=≈]\s*[\d.,]+\s*[a-zA-Z]', stripped)
                and len(stripped) < 150
                and stripped.count("=") >= 1):
            continue
        # Skip lines starting with bullet/list markers
        if re.match(r'^[-•*–·]\s|^\d+[.)]', stripped):
            continue
        # Skip short fragments (<5 words)
        if len(stripped.split()) < 5:
            continue
        lines.append(stripped)

    prose = " ".join(l for l in lines if l).strip()
    # Collapse extra spaces
    prose = re.sub(r'\s{2,}', ' ', prose)

    # Split into sentences and take max_sentences
    sentences = re.split(r'(?<=[.!?])\s+', prose)
    result = " ".join(sentences[:max_sentences])

    # Ensure it ends with a period
    if result and not result[-1] in ".!?":
        result += "."
    return result


def _parse_sections(text: str) -> List[tuple]:
    """
    Parse final_report text into (title, body) tuples.
    Recognizes: ## Heading, ### Heading, 1. HEADING, ALL CAPS headings.
    Pre-processes: joins CROSS-DOMAIN FLAG continuation lines.
    """
    # Pre-process: join CROSS-DOMAIN FLAG line with its continuation
    lines_raw = text.splitlines()
    joined = []
    j = 0
    while j < len(lines_raw):
        line = lines_raw[j]
        stripped = line.strip()
        if "CROSS-DOMAIN FLAG" in stripped.upper() and not stripped.endswith("."):
            # Collect continuation lines
            combined = stripped
            j += 1
            while j < len(lines_raw):
                nxt = lines_raw[j].strip()
                if (not nxt
                        or re.match(r'^#{1,4}\s', nxt)
                        or re.match(r'^\d+\.\s+[A-Z]', nxt)
                        or re.match(r'^[A-Z][A-Z &/:0-9-]{4,}$', nxt)):
                    break
                combined = combined + " " + nxt
                j += 1
            joined.append(combined)
        else:
            joined.append(line)
            j += 1
    text = "\n".join(joined)

    heading_re = re.compile(
        r'^(?:'
        r'#{1,4}\s+.{3,}'
        r'|(?:\d+\.?\d*\.?\s+[A-Z].{2,})'
        r'|(?:\*\*[A-Z].{3,}\*\*)'
        r'|(?:[A-Z][A-Z &/:0-9\-]{4,})$'
        r')'
    )
    sections, cur_title, cur_lines = [], "", []
    for line in text.splitlines():
        s = line.strip()
        if s and heading_re.match(s):
            body = "\n".join(cur_lines).strip()
            if body:
                sections.append((cur_title, body))
            cur_title = re.sub(r'^[#\s*0-9.]+', '', s).strip().strip("*:").strip()
            cur_lines = []
        else:
            cur_lines.append(line)
    body = "\n".join(cur_lines).strip()
    if body:
        sections.append((cur_title, body))
    return sections


def _clean_heading(text: str) -> str:
    """Strip markdown heading markers from a line."""
    return re.sub(r'^#{1,6}\s*', '', text).strip().strip("*:").strip()


def _is_md_heading(text: str) -> bool:
    return bool(re.match(r'^#{1,6}\s+.{2,}', text))


def _is_section_divider(text: str) -> bool:
    return bool(re.match(r'^[=\-]{4,}\s*$', text))


def _render_body(doc, text: str, base_size: float = SZ_BODY):
    """
    Render free-form agent output into the document with academic typography.
    Handles: markdown headings, bullets, numbered lists, equations, tables,
    CRITICAL/HIGH/MEDIUM priority lines, CROSS-DOMAIN FLAGs, inline bold.
    """
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i].strip()
        i += 1

        if not raw:
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(3)
            continue

        # Horizontal rule → skip
        if _is_section_divider(raw):
            continue

        # Markdown heading (## / ###) → bold italic inline sub-heading
        if _is_md_heading(raw):
            heading_text = _clean_heading(raw)
            if not heading_text:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.keep_with_next = True
            _run(p, heading_text, bold=True, italic=True,
                 size=base_size, color=C_NAVY)
            continue

        # ALL-CAPS standalone heading (e.g. "MATERIALS ANALYSIS")
        if (re.match(r'^[A-Z][A-Z &/:0-9-]{4,}$', raw)
                and not re.search(r'[=<>]', raw)
                and len(raw) < 60):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.keep_with_next = True
            _run(p, raw.title(), bold=True, italic=True,
                 size=base_size, color=C_NAVY)
            continue

        # Markdown table row
        if raw.startswith("|") and raw.count("|") >= 2:
            if re.match(r'^[|\-\s:]+$', raw):
                continue
            cells = [c.strip() for c in raw.strip("|").split("|")]
            if all(len(c) <= 2 for c in cells):
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.4)
            _run(p, "   |   ".join(cells), size=base_size - 1.5,
                 color=C_GREY, font=F_MONO)
            continue

        # Bullet point
        if re.match(r'^[-•*–·]\s', raw):
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Cm(0.7)
            p.paragraph_format.first_line_indent = Cm(-0.45)
            _run(p, "•  ", bold=True, size=base_size, color=C_NAVY)
            _run(p, raw.lstrip("-•*–· ").strip(), size=base_size, color=C_BLACK)
            _set_line_spacing(p, 1.1)
            continue

        # Numbered list (1. text or 1) text)
        m_num = re.match(r'^(\d+[.)]\s+)(.*)', raw)
        if m_num:
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Cm(0.7)
            p.paragraph_format.first_line_indent = Cm(-0.45)
            _run(p, m_num.group(1).rstrip() + "  ", bold=True,
                 size=base_size, color=C_NAVY)
            _run(p, m_num.group(2).strip(), size=base_size, color=C_BLACK)
            _set_line_spacing(p, 1.1)
            continue

        # CRITICAL / HIGH / MEDIUM / LOW priority lines
        m_prio = re.match(r'^(CRITICAL|HIGH|MEDIUM|LOW):\s*(.*)', raw, re.IGNORECASE)
        if m_prio:
            level     = m_prio.group(1).upper()
            body_text = m_prio.group(2).strip()
            c_map = {"CRITICAL": C_ERR, "HIGH": C_WARN, "MEDIUM": C_NAVY, "LOW": C_MUTED}
            c = c_map.get(level, C_NAVY)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.4)
            _run(p, f"[{level}]  ", bold=True, size=base_size - 0.5, color=c)
            _run(p, body_text, size=base_size, color=C_BLACK)
            _set_line_spacing(p, 1.1)
            continue

        # Cross-domain flag
        if "CROSS-DOMAIN FLAG" in raw.upper():
            body_clean = re.sub(r'CROSS[-\s]?DOMAIN\s+FLAG[:\s]*',
                                '', raw, flags=re.IGNORECASE).strip()
            _note_box(doc, "⚑  Cross-Domain Flag", body_clean)
            continue

        # Equation / calculation line (= number unit pattern)
        if (re.search(r'[=≈]\s*[\d.,]+\s*[a-zA-Z]', raw)
                and not raw.endswith(":")
                and len(raw) < 220):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.left_indent  = Cm(1.4)
            _run(p, raw, size=base_size - 1, color=C_BLACK, font=F_MONO)
            continue

        # Regular body paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _set_line_spacing(p, 1.15)
        parts = re.split(r'(\*\*[^*]+\*\*)', raw)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                _run(p, part[2:-2], bold=True, size=base_size, color=C_BLACK)
            else:
                _run(p, part, size=base_size, color=C_BLACK)


# ═══════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════

def _build_cover(doc, brief: str, domains: list, mode: int,
                 total_cost: float, kur: float, round_scores: list,
                 agent_log: list, mode_label: str, doc_no: str):
    """Academic-style cover page."""

    # Vertical spacers
    for _ in range(5):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    # System label
    sl = doc.add_paragraph()
    sl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(sl, "ENGINEERING AI  ·  MULTI-AGENT ANALYSIS SYSTEM",
         size=8, color=C_MUTED)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

    # Document type
    dt = doc.add_paragraph()
    dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(dt, "TECHNICAL ANALYSIS REPORT", bold=True, size=20, color=C_NAVY)
    dt.paragraph_format.space_after = Pt(4)

    # Top rule
    _h_rule(doc, C_NAVY_HX, 14)

    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(14)

    # Project title (brief excerpt)
    brief_short = brief[:160].replace("\n", " ").strip()
    if len(brief) > 160:
        brief_short += "..."

    pt_label = doc.add_paragraph()
    pt_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(pt_label, "ANALYSIS SUBJECT", size=8, color=C_MUTED)
    pt_label.paragraph_format.space_after = Pt(3)

    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(pt, brief_short, size=11, color=C_BLACK)
    _set_line_spacing(pt, 1.2)
    pt.paragraph_format.space_after = Pt(20)

    # Metadata table (2-column, no borders)
    mode_labels = {1: "Mode 1 — Single Agent", 2: "Mode 2 — Dual Agent",
                   3: "Mode 3 — Semi-Automatic", 4: "Mode 4 — Full Automatic"}
    best_score  = max((r.get("puan", 0) for r in (round_scores or [])), default=None)
    n_domains   = len(domains)
    n_agents    = len(agent_log) if agent_log else 0
    n_rounds    = len(round_scores) if round_scores else 1

    meta_rows = [
        ("Document No.",     doc_no),
        ("Date",             datetime.datetime.now().strftime("%B %d, %Y")),
        ("Analysis Mode",    mode_labels.get(mode, mode_label)),
        ("Engineering Domains", ", ".join(domains) if domains else "—"),
        ("Analysis Rounds",  str(n_rounds)),
        ("Total Agents",     str(n_agents)),
        ("Quality Score",    f"{best_score}/100" if best_score else "—"),
        ("Total API Cost",   f"${total_cost:.4f} USD   (≈ {total_cost * kur:.2f} TL)"),
    ]

    mt = doc.add_table(rows=len(meta_rows), cols=2)
    _no_border_table(mt)
    mt.columns[0].width = Cm(4.5)
    mt.columns[1].width = Cm(11.5)

    for ri, (label, value) in enumerate(meta_rows):
        bg = "F5F7FA" if ri % 2 == 0 else "FAFBFC"
        lc = mt.rows[ri].cells[0]
        vc = mt.rows[ri].cells[1]
        for cell, txt, bold in ((lc, label, True), (vc, value, False)):
            _shade(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.left_indent  = Cm(0.3)
            _run(p, txt, bold=bold, size=9.5,
                 color=C_NAVY if bold else C_BLACK)

    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(10)

    # Bottom rule
    _h_rule(doc, C_NAVY_HX, 8)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(8)

    # Footer note on cover
    fn = doc.add_paragraph()
    fn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(fn, "This document was generated automatically by the Engineering AI multi-agent analysis system. "
         "Results should be reviewed by a qualified engineer before use in design decisions.",
         italic=True, size=8.5, color=C_MUTED)

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# SECTION 1 — INTRODUCTION
# ═══════════════════════════════════════════════════════════════

def _build_introduction(doc, brief: str, domains: list, mode: int,
                        assumptions: list = None):
    _h1(doc, "1.", "Introduction")

    _h2(doc, "1.1", "Problem Statement")
    _body(doc, brief.strip())

    _h2(doc, "1.2", "Scope and Objectives")
    scope_text = (
        f"This analysis addresses the engineering problem stated above by deploying "
        f"{len(domains)} specialist domain agent(s): {', '.join(domains)}. "
        f"The analysis was conducted in accordance with Mode {mode} of the "
        f"Engineering AI multi-agent framework. "
        f"Each domain agent provides independent analysis from either a theoretical "
        f"(Expert A) or applied (Expert B) perspective, followed by cross-validation, "
        f"observer evaluation, and synthesis into this report."
    )
    _body(doc, scope_text)

    _h2(doc, "1.3", "Assumptions and Limitations")
    if assumptions:
        for a in assumptions:
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Cm(0.6)
            p.paragraph_format.first_line_indent = Cm(-0.4)
            _run(p, "•  ", bold=True, size=SZ_BODY, color=C_NAVY)
            _run(p, a, size=SZ_BODY, color=C_BLACK)
    else:
        # Try to extract [ASSUMPTION] labeled lines from brief context
        _body(doc, (
            "Assumptions embedded within the analysis are explicitly labeled "
            "[ASSUMPTION] within the Technical Findings section. Conservative "
            "engineering estimates have been applied where specific input data "
            "was unavailable, and each such estimate is documented with its "
            "classification (standard simplification, problem-specific, or "
            "conservative bound) and its quantitative impact on the result. "
            "This analysis represents a preliminary multi-perspective assessment "
            "and does not substitute for detailed design calculations, "
            "site-specific measurements, or physical testing."
        ))


# ═══════════════════════════════════════════════════════════════
# SECTION 2 — METHODOLOGY
# ═══════════════════════════════════════════════════════════════

def _build_methodology(doc, mode: int, domains: list, max_rounds: int,
                       round_scores: list):
    _h1(doc, "2.", "Methodology")

    _h2(doc, "2.1", "Analysis Framework")
    mode_descs = {
        1: {
            "label": "Single-Agent Analysis (Mode 1)",
            "body": (
                f"The analysis employed Mode 1 (Single-Agent), deploying one specialist "
                f"expert per engineering domain. Each domain agent performed independent "
                f"analysis, producing quantitative findings, safety assessments, and "
                f"domain-specific recommendations. This mode is optimised for rapid "
                f"preliminary assessment where a single technical perspective per domain "
                f"is sufficient. A total of {len(domains)} domain(s) were activated: "
                f"{', '.join(domains)}."
            )
        },
        2: {
            "label": "Dual-Agent Analysis (Mode 2)",
            "body": (
                f"The analysis employed Mode 2 (Dual-Agent), deploying two specialist "
                f"experts per engineering domain — Expert A providing theoretical "
                f"analysis and Expert B providing applied, field-validated perspective. "
                f"Positions were compared and conflicts resolved by a dedicated "
                f"Conflict Resolution agent. This mode provides a theory-versus-practice "
                f"dialogue that surfaces design tensions not visible from a single "
                f"perspective. Active domains: {', '.join(domains)}."
            )
        },
        3: {
            "label": "Semi-Automatic Full-Loop Analysis (Mode 3)",
            "body": (
                f"The analysis employed Mode 3 (Semi-Automatic), combining user-supplied "
                f"clarifications with a full multi-round feedback loop. The Prompt "
                f"Engineer agent first identified missing parameters and presented them "
                f"to the user; responses were used to strengthen the problem brief "
                f"before domain analysis commenced. Dual-perspective domain agents "
                f"then executed across {len(domains)} domain(s) over up to "
                f"{max_rounds} rounds, with Observer-directed iteration. "
                f"Active domains: {', '.join(domains)}."
            )
        },
        4: {
            "label": "Full-Automatic Deep Analysis (Mode 4)",
            "body": (
                f"The analysis employed Mode 4 (Full-Automatic), the highest-fidelity "
                f"mode. The Prompt Engineer agent autonomously identified and resolved "
                f"missing parameters before deploying dual-perspective domain agents "
                f"across {len(domains)} domain(s). Up to {max_rounds} iterative "
                f"analysis rounds were executed, with the Observer meta-agent providing "
                f"quality-directed feedback between rounds. Post-loop synthesis, "
                f"alternative scenario, cost, standards, and simulation agents "
                f"supplemented the domain findings. Active domains: "
                f"{', '.join(domains)}."
            )
        },
    }
    m_info = mode_descs.get(mode, {
        "label": f"Mode {mode}",
        "body": f"Analysis mode {mode} was applied across {len(domains)} domain(s): {', '.join(domains)}."
    })
    _body(doc, m_info["body"])

    _h2(doc, "2.2", "Quality Assurance")
    if round_scores:
        scores = [r.get("puan", 0) for r in round_scores]
        score_str = ", ".join(f"Round {r['tur']}: {r.get('puan','—')}/100"
                              for r in round_scores)
        _body(doc, (
            f"Each analysis round was evaluated by an independent Observer meta-agent "
            f"using a weighted rubric: technical accuracy (30%), internal consistency "
            f"(25%), assumption transparency (20%), analysis depth (15%), and "
            f"cross-validation quality (10%). Scores this analysis: {score_str}. "
            f"A score of 85/100 or above triggers early termination."
        ))
    else:
        _body(doc, (
            "Quality control was performed by an independent Observer meta-agent "
            "evaluating technical accuracy, internal consistency, assumption transparency, "
            "analysis depth, and cross-agent consistency."
        ))

    _h2(doc, "2.3", "Domain Coverage")
    rows = [[str(i+1), d] for i, d in enumerate(domains)]
    _make_table(doc,
                ["No.", "Engineering Domain"],
                rows,
                [1.0, 13.0],
                caption="Table 2.1.  Engineering domains active in this analysis.")


# ═══════════════════════════════════════════════════════════════
# SECTION 3 — TECHNICAL FINDINGS
# ═══════════════════════════════════════════════════════════════

def _extract_section(text: str, markers: list, stop_markers: list = None) -> str:
    """
    Extract a named section from report text.
    markers: list of section name patterns to search for
    stop_markers: list of patterns that terminate the section
    Returns cleaned body text (no heading line included).
    """
    stop_pat = stop_markers or []
    # Build stop pattern: next heading (## or numbered) or specific markers
    stop_group = (
        r'(?=\n#{1,4}\s|\n\d+\.\s+[A-Z]|\n[A-Z]{4,}\n'
        + (r'|\n' + r'|\n'.join(re.escape(s) for s in stop_pat) if stop_pat else r'')
        + r'|\Z)'
    )
    for marker in markers:
        pattern = rf'(?:^|\n){marker}[^\n]*\n+(.*?)' + stop_group
        m = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if m and len(m.group(1).strip()) > 30:
            return m.group(1).strip()
    return ""


def _build_findings(doc, final_report: str, domains: list):
    _h1(doc, "3.", "Technical Findings")

    # Skip these sections — they belong to Discussion / Conclusions / References
    SKIP_RE = re.compile(
        r'executive summary|abstract|özet|yönetici|'
        r'conclusion|sonuç|recommendation|öneri|'
        r'next step|reference|kaynaklar?|appendix|ek[\s:]|methodology',
        re.IGNORECASE
    )
    # Note: cross-domain, risk assessment are kept in findings (NOT skipped)
    # Discussion (Section 4) generates its own narrative — no raw re-extraction

    sections = _parse_sections(final_report)
    findings_sections = [
        (t, b) for (t, b) in sections
        if t and not SKIP_RE.search(t)
    ]

    if not findings_sections:
        _render_body(doc, final_report)
        return

    sub_counter = 1
    for title, body in findings_sections:
        _h2(doc, f"3.{sub_counter}", title)
        sub_counter += 1
        _render_body(doc, body)
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(4)


# ═══════════════════════════════════════════════════════════════
# SECTION 4 — DISCUSSION
# ═══════════════════════════════════════════════════════════════

def _build_discussion(doc, final_report: str, round_scores: list):
    """
    Section 4: Discussion — synthesised commentary and quality assessment.
    Does NOT re-extract raw domain content (that belongs in Section 3).
    Generates narrative from Observer scores and cross-validation findings.
    """
    _h1(doc, "4.", "Discussion")

    # 4.1 Cross-domain synthesis — narrative paragraph (not raw extraction)
    _h2(doc, "4.1", "Cross-Domain Synthesis")
    # Count how many domains mentioned in report
    domain_mentions = len(re.findall(
        r'(?:EXPERT|Engineer|Analyst|Analysis)[^\n]{0,40}(?=\n)',
        final_report, re.IGNORECASE
    ))
    _body(doc, (
        "The domain analyses presented in Section 3 were subjected to numerical "
        "cross-validation to verify dimensional consistency, order-of-magnitude "
        "plausibility, and inter-agent coherence. Material properties, load "
        "parameters, and derived safety factors were checked across all active "
        "domains. Where agents raised cross-domain flags — indicating a finding "
        "in one discipline with consequences for another — these are annotated "
        "within the relevant subsection and must be resolved jointly before "
        "design can advance."
    ))

    # 4.2 FMEA chart (visual only — text summary already in findings)
    rpn_img = chart_rpn(final_report)
    if rpn_img:
        _h2(doc, "4.2", "Risk Priority Summary")
        _body(doc, (
            "Failure modes identified across the domain analyses were ranked "
            "by Risk Priority Number (RPN = Severity × Occurrence × "
            "Detectability). The chart below provides a comparative view. "
            "Items with RPN ≥ 200 are classified Critical and require "
            "immediate design action; items in the range 100–199 are High "
            "priority; items below 100 are Medium priority."
        ))
        _embed_chart(doc, rpn_img,
                     "Figure 4.1.  FMEA Risk Priority Numbers. "
                     "Critical: RPN ≥ 200  |  High: 100–199  |  Medium: < 100.")

    # 4.3 Observer quality assessment
    if round_scores:
        _h2(doc, "4.3", "Quality Assessment")
        final_score = round_scores[-1].get("puan") if round_scores else None
        n_rounds    = len(round_scores)
        if final_score is not None:
            score_label = (
                "satisfactory (≥ 85/100)" if final_score >= 85 else
                "acceptable (70–84/100)"  if final_score >= 70 else
                "below target (< 70/100)"
            )
            rounds_str = (
                f"across {n_rounds} analysis rounds"
                if n_rounds > 1 else "in a single analysis round"
            )
            # Build score progression narrative
            if len(round_scores) > 1:
                progression = ", ".join(
                    f"Round {r['tur']}: {r.get('puan','—')}/100"
                    for r in round_scores
                )
                prog_text = (
                    f" Score progression: {progression}. "
                    f"The {'improvement' if round_scores[-1].get('puan',0) > round_scores[0].get('puan',0) else 'consistency'} "
                    f"across rounds reflects the iterative feedback mechanism of the multi-agent pipeline."
                )
            else:
                prog_text = ""

            _body(doc, (
                f"The independent Observer meta-agent evaluated all domain outputs "
                f"{rounds_str} using a five-criterion weighted rubric: "
                f"technical accuracy (30%), internal consistency (25%), "
                f"assumption transparency (20%), analysis depth (15%), and "
                f"cross-validation quality (10%). "
                f"The final quality score is {final_score}/100 — {score_label}."
                f"{prog_text}"
            ))

            # Quality score table for multi-round
            if len(round_scores) > 1:
                rows = []
                for r in round_scores:
                    p = r.get("puan", 0)
                    status = (
                        ("✓ Target met",    C_OK,   True)  if p >= 85 else
                        ("~  Acceptable",   C_WARN, False) if p >= 70 else
                        ("✗ Below target",  C_ERR,  True)
                    )
                    rows.append([f"Round {r['tur']}", f"{p} / 100", status])
                _make_table(doc, ["Round", "Quality Score", "Status"],
                            rows, [2.5, 3.5, 8.0],
                            caption="Table 4.1.  Observer quality scores by analysis round.")


# ═══════════════════════════════════════════════════════════════
# SECTION 5 — CONCLUSIONS & RECOMMENDATIONS
# ═══════════════════════════════════════════════════════════════

def _build_conclusions(doc, final_report: str):
    _h1(doc, "5.", "Conclusions and Recommendations")

    # 5.1 Conclusions — stop BEFORE recommendations and references
    conc_text = _extract_section(
        final_report,
        [r"CONCLUSION", r"SONUÇ", r"KEY FINDINGS?"],
        stop_markers=["RECOMMENDATION", "NEXT STEP", "ÖNERI", "REFERENCE",
                      "KAYNAKLAR", "SUMMARY", "RISK ASSESSMENT"]
    )
    # Strip equations and calculation lines from conclusions text
    if conc_text:
        clean_lines = []
        for ln in conc_text.splitlines():
            stripped = ln.strip()
            # Skip pure equation lines
            if (re.search(r'[=≈]\s*[\d.,]+\s*[a-zA-Z/·]', stripped)
                    and stripped.count("=") >= 1
                    and len(stripped) < 200):
                continue
            clean_lines.append(ln)
        conc_text = "\n".join(clean_lines).strip()
    _h2(doc, "5.1", "Conclusions")
    if conc_text:
        _render_body(doc, conc_text)
    else:
        # Extract key findings bullet lines from anywhere in final_report
        finding_lines = re.findall(
            r'(?:^|\n)(?:\d+\.\s+|[-•]\s*)(.{40,180})(?=\n|$)',
            _strip_markup(final_report)
        )
        # Filter out headings and non-substantive lines
        finding_lines = [
            l.strip() for l in finding_lines
            if not re.match(r'^[A-Z\s]{4,}$', l.strip())
            and len(l.strip().split()) > 8
            and not l.strip().startswith(("CRITICAL", "HIGH", "MEDIUM", "LOW"))
        ][:5]

        # Filter out raw calculation/equation lines — keep interpretive conclusions
        finding_lines = [
            l for l in finding_lines
            if not (re.search(r'[=≈]\s*[\d.,]+\s*[a-zA-Z]', l) and l.count("=") >= 1)
            and not re.match(r'^[A-Z]\s*[=:]\s*', l)
            and len(l.split()) >= 8
        ][:5]

        if finding_lines:
            _body(doc, (
                "The following principal conclusions are drawn from the "
                "domain analyses presented in Section 3:"
            ))
            for line in finding_lines:
                p = doc.add_paragraph()
                p.paragraph_format.space_after  = Pt(3)
                p.paragraph_format.left_indent  = Cm(0.7)
                p.paragraph_format.first_line_indent = Cm(-0.45)
                _run(p, "•  ", bold=True, size=SZ_BODY, color=C_NAVY)
                _run(p, line.lstrip("-•*0123456789.) "), size=SZ_BODY, color=C_BLACK)
                _set_line_spacing(p, 1.1)
        else:
            _body(doc, (
                "The principal technical conclusions of this analysis are "
                "presented in the domain-specific subsections of Section 3. "
                "Quantitative results, safety factors, and failure assessments "
                "constitute the analytical basis for the recommendations below."
            ))

    # 5.2 Recommendations — stop BEFORE references
    rec_text = _extract_section(
        final_report,
        [r"RECOMMENDATION", r"NEXT STEP", r"ÖNERİ"],
        stop_markers=["REFERENCE", "KAYNAKLAR", "APPENDIX"]
    )
    _h2(doc, "5.2", "Recommendations")
    if rec_text:
        _render_body(doc, rec_text)
    else:
        # Fall back to extracting inline CRITICAL/HIGH/MEDIUM lines
        prio_lines = re.findall(
            r'^(CRITICAL|HIGH|MEDIUM|LOW):\s*(.+)$',
            final_report, re.MULTILINE | re.IGNORECASE
        )
        if prio_lines:
            for level, body_text in prio_lines:
                c_map = {"CRITICAL": C_ERR, "HIGH": C_WARN, "MEDIUM": C_NAVY, "LOW": C_MUTED}
                c = c_map.get(level.upper(), C_NAVY)
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Cm(0.4)
                _run(p, f"[{level.upper()}]  ", bold=True, size=SZ_BODY - 0.5, color=c)
                _run(p, body_text.strip(), size=SZ_BODY, color=C_BLACK)
                _set_line_spacing(p, 1.1)
        else:
            _body(doc, (
                "Specific recommendations are identified within the Technical "
                "Findings section. All recommendations are classified by priority: "
                "CRITICAL — must be addressed before design can advance; "
                "HIGH — should be resolved in the next design iteration; "
                "MEDIUM — address during the detailed design phase."
            ))


# ═══════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════

def _build_references(doc, final_report: str):
    _h1(doc, "", "References")

    ref_text = _extract_section(
        final_report,
        [r"REFERENCES?", r"KAYNAKLAR?"],
        stop_markers=["APPENDIX"]
    )
    if ref_text:
        # Render each reference line as a hanging-indent numbered entry
        lines = [l.strip() for l in ref_text.splitlines() if l.strip()]
        for j, line in enumerate(lines, 1):
            # If not already numbered, add number
            if not re.match(r'^\[?\d+\]?', line):
                line = f"[{j}]  {line}"
            p = doc.add_paragraph()
            p.paragraph_format.space_after        = Pt(4)
            p.paragraph_format.left_indent        = Cm(0.9)
            p.paragraph_format.first_line_indent  = Cm(-0.6)
            _set_line_spacing(p, 1.1)
            _run(p, line, size=SZ_SMALL, color=C_BLACK)
    else:
        _body(doc, (
            "References cited by the analysis agents are documented "
            "in-text within the Technical Findings section. Standard sources "
            "include ASM Handbooks, ASME/ASTM/MIL-SPEC standards, Eurocode "
            "series, and peer-reviewed engineering literature."
        ), spacing=1.0)


# ═══════════════════════════════════════════════════════════════
# APPENDIX A — ANALYSIS METADATA
# ═══════════════════════════════════════════════════════════════

def _build_appendix_a(doc, domains: list, mode: int, round_scores: list,
                      total_cost: float, kur: float, agent_log: list,
                      max_rounds: int):
    _appendix_h1(doc, "A", "Analysis Metadata")

    _h2(doc, "A.1", "Session Summary")

    mode_labels = {1: "Single Agent", 2: "Dual Agent",
                   3: "Semi-Automatic", 4: "Full Automatic"}
    n_domain_agents  = sum(1 for a in (agent_log or [])
                           if a.get("key","").endswith(("_a","_b")))
    n_support_agents = len(agent_log or []) - n_domain_agents

    rows = [
        ["Analysis Mode",        f"{mode}  —  {mode_labels.get(mode,'—')}"],
        ["Active Domains",       str(len(domains))],
        ["Rounds Completed",     str(len(round_scores)) if round_scores else "1"],
        ["Domain Agents",        str(n_domain_agents)],
        ["Support Agents",       str(n_support_agents)],
        ["Total Agents",         str(len(agent_log) if agent_log else 0)],
        ["Total Input Tokens",   "—"],
        ["Total Output Tokens",  "—"],
        ["Total API Cost (USD)", f"${total_cost:.5f}"],
        ["Approximate Cost (TL)", f"≈ {total_cost * kur:.2f} TL  (exchange rate: {kur:.1f})"],
    ]
    _make_table(doc, ["Parameter", "Value"], rows, [5.0, 10.0],
                caption="Table A.1.  Analysis session parameters and resource utilization.")

    # Quality chart
    if round_scores:
        _h2(doc, "A.2", "Quality Score History")
        img = chart_quality_scores(round_scores)
        if img:
            _embed_chart(doc, img,
                         "Figure A.1.  Observer quality score per analysis round. "
                         "Target threshold: 85/100. Score below 70/100 triggers mandatory revision.")
        else:
            score_rows = [
                [f"Round {r['tur']}", str(r.get('puan','—')) + " / 100",
                 ("✓ Target met", C_OK, True)  if (r.get('puan') or 0) >= 85 else
                 ("~  Acceptable",C_WARN,False) if (r.get('puan') or 0) >= 70 else
                 ("✗ Below target",C_ERR,True)]
                for r in round_scores
            ]
            _make_table(doc, ["Round", "Quality Score", "Status"],
                        score_rows, [2.5, 3.5, 8.0],
                        caption="Table A.2.  Observer quality scores by round.")

    # Cost chart
    if agent_log:
        _h2(doc, "A.3", "Agent Cost Distribution")
        img = chart_agent_cost(agent_log)
        if img:
            _embed_chart(doc, img,
                         f"Figure A.2.  API cost per agent. Total: ${total_cost:.4f} USD "
                         f"(≈ {total_cost*kur:.2f} TL).")
        else:
            cost_rows = sorted(
                [[a.get("name","?"), f"${a.get('cost',0):.5f}"]
                 for a in agent_log if a.get("cost",0) > 0],
                key=lambda x: float(x[1][1:]), reverse=True
            )[:15]
            _make_table(doc, ["Agent", "Cost (USD)"], cost_rows, [11.0, 4.0],
                        caption="Table A.3.  Per-agent API cost.")


# ═══════════════════════════════════════════════════════════════
# APPENDIX B — AGENT ACTIVITY SUMMARY
# ═══════════════════════════════════════════════════════════════

def _summarise_output(text: str, max_words: int = 150) -> str:
    """Condense agent output to key sentences — first substantive content."""
    if not text or text.strip().startswith("ERROR"):
        return "(No output recorded.)"
    # Strip headers / blank lines
    lines = [l.strip() for l in text.splitlines() if l.strip()
             and not re.match(r'^#{1,4}\s|^[=\-]{5,}$', l.strip())]
    joined = " ".join(lines)
    words  = joined.split()
    if len(words) <= max_words:
        return joined
    return " ".join(words[:max_words]) + " [...]"


def _build_appendix_b(doc, agent_log: list):
    if not agent_log:
        return

    _appendix_h1(doc, "B", "Agent Activity Summary")

    _body(doc, (
        "This appendix documents the outputs of all agents that participated in "
        "the analysis. Domain agent outputs are presented in full (condensed to "
        "preserve key findings); observer, cross-validation, and risk agent "
        "outputs are presented in full to allow independent review of the "
        "quality control process."
    ))

    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    # Classify agents
    try:
        import sys as _sys, os as _os
        _sys.path.insert(0, _os.path.join(_os.path.dirname(__file__)))
        from config.agents_config import AGENTS as _DA
    except Exception:
        _DA = {}

    # Priority agents shown in full (expanded summary)
    FULL_TEXT_KEYS = {
        "gozlemci", "capraz_dogrulama", "varsayim_belirsizlik",
        "risk_guvenilirlik", "celisiki_cozum", "soru_uretici",
        "literatur_patent"
    }

    def _model_tag(model_str):
        if "opus"   in (model_str or ""): return "Opus"
        if "sonnet" in (model_str or ""): return "Sonnet"
        if "haiku"  in (model_str or ""): return "Haiku"
        return model_str or "—"

    domain_log   = [a for a in agent_log if a.get("key","") in _DA]
    qa_log       = [a for a in agent_log
                    if a.get("key","") in FULL_TEXT_KEYS]
    other_log    = [a for a in agent_log
                    if a.get("key","") not in _DA
                    and a.get("key","") not in FULL_TEXT_KEYS]

    # B.1 Domain agents — condensed table
    if domain_log:
        _h2(doc, "B.1", "Domain Agent Outputs")
        _body(doc, (
            "Each domain agent conducted independent analysis from a specialist "
            "perspective. Outputs are condensed to key findings and numerical results."
        ))
        d_rows = [
            [a.get("name","?"),
             _model_tag(a.get("model","")),
             f"${a.get('cost',0):.5f}",
             _summarise_output(a.get("output",""), 250)]
            for a in domain_log
        ]
        _make_table(doc,
                    ["Agent", "Model", "Cost (USD)", "Key Findings (condensed)"],
                    d_rows, [3.5, 1.5, 2.0, 9.0],
                    caption="Table B.1.  Domain agent outputs — condensed to principal findings.")

    # B.2 QA agents — fuller text per agent
    if qa_log:
        _h2(doc, "B.2", "Quality Assurance Agent Outputs")
        _body(doc, (
            "The following agents performed quality control, cross-validation, "
            "risk assessment, and critical questioning. Their outputs are "
            "presented in greater detail to allow independent review."
        ))
        for j, a in enumerate(qa_log, 1):
            name    = a.get("name","?")
            model   = _model_tag(a.get("model",""))
            cost    = a.get("cost",0)
            output  = (a.get("output") or "").strip()
            clean   = _strip_markup(output)
            # Show up to 400 words for QA agents
            words   = clean.split()
            if len(words) > 400:
                clean = " ".join(words[:400]) + " [condensed]"

            _h3(doc, f"B.2.{j}  {name}  ·  {model}  ·  ${cost:.5f}")
            if clean:
                _render_body(doc, clean, base_size=SZ_SMALL)
            sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    # B.3 Remaining support agents — brief table
    if other_log:
        _h2(doc, "B.3", "Other Support Agent Outputs")
        o_rows = [
            [a.get("name","?"),
             _model_tag(a.get("model","")),
             f"${a.get('cost',0):.5f}",
             _summarise_output(a.get("output",""), 80)]
            for a in other_log
        ]
        _make_table(doc,
                    ["Agent", "Model", "Cost (USD)", "Summary"],
                    o_rows, [4.0, 1.5, 2.0, 8.5],
                    caption="Table B.3.  Other support and synthesis agent outputs.")


# ═══════════════════════════════════════════════════════════════
# HEADER & FOOTER
# ═══════════════════════════════════════════════════════════════

def _setup_running_headers(doc, brief_short: str, doc_no: str):
    date_str = datetime.datetime.now().strftime("%Y-%m-%d")
    for section in doc.sections:
        # Header: left = system name, right = doc number
        hp = section.header.paragraphs[0]
        hp.clear()
        hp.paragraph_format.space_after = Pt(2)
        _run(hp, "Engineering AI  ·  Technical Analysis Report", size=8, color=C_MUTED)
        # Tab stop to right-align doc number
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Footer: left = brief excerpt, right = date
        fp = section.footer.paragraphs[0]
        fp.clear()
        _run(fp, f"{brief_short[:60]}{'...' if len(brief_short)>60 else ''}",
             size=7.5, color=C_MUTED, italic=True)
        _run(fp, f"    {date_str}    {doc_no}", size=7.5, color=C_MUTED)
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT


# ═══════════════════════════════════════════════════════════════
# MAIN FUNCTION
# ═══════════════════════════════════════════════════════════════

def generate_docx_report(
    brief:        str,
    final_report: str,
    domains:      List[str],
    round_scores: List[Dict] = None,
    agent_log:    List[Dict] = None,
    total_cost:   float      = 0.0,
    kur:          float      = 44.0,
    mode:         int        = 4,
    max_rounds:   int        = 3,
) -> bytes:

    round_scores = round_scores or []
    agent_log    = agent_log    or []

    # Document number
    doc_no = f"EAI-{datetime.datetime.now().strftime('%Y%m%d-%H%M')}-M{mode}"

    mode_labels = {1: "Single Agent", 2: "Dual Agent",
                   3: "Semi-Automatic", 4: "Full Automatic"}
    mode_label  = mode_labels.get(mode, "Custom")
    brief_short = brief[:90].replace("\n", " ").strip()

    # ── Document setup ───────────────────────────────────────
    doc = Document()
    _set_margins(doc, top=2.5, right=2.0, bottom=2.5, left=2.5)
    _setup_styles(doc)

    # ── 1. Cover page ────────────────────────────────────────
    _build_cover(doc, brief, domains, mode, total_cost, kur,
                 round_scores, agent_log, mode_label, doc_no)

    # ── 2. Abstract ──────────────────────────────────────────
    _h1(doc, "", "Abstract")
    abstract = _extract_abstract(final_report)
    if abstract:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.8)
        p.paragraph_format.right_indent = Cm(0.8)
        p.paragraph_format.space_after  = Pt(8)
        _set_line_spacing(p, 1.15)
        _run(p, abstract, italic=True, size=SZ_ABSTRACT, color=C_BLACK)
    else:
        _body(doc, f"This report presents the results of a multi-agent engineering "
              f"analysis of the following problem: {brief_short}. "
              f"The analysis was conducted in {mode_label} and covers "
              f"{len(domains)} engineering domain(s): {', '.join(domains)}.")

    # Keyword line
    kw = doc.add_paragraph()
    kw.paragraph_format.left_indent = Cm(0.8)
    kw.paragraph_format.space_after = Pt(0)
    _run(kw, "Keywords:  ", bold=True, size=SZ_SMALL, color=C_NAVY)
    _run(kw, ", ".join(domains) + f"; multi-agent analysis; engineering AI; {mode_label.lower()}",
         italic=True, size=SZ_SMALL, color=C_GREY)

    doc.add_page_break()

    # ── 3. Introduction ──────────────────────────────────────
    _build_introduction(doc, brief, domains, mode)

    # ── 4. Methodology ───────────────────────────────────────
    _build_methodology(doc, mode, domains, max_rounds, round_scores)

    # ── 5. Technical Findings ────────────────────────────────
    _build_findings(doc, final_report, domains)

    # ── 6. Discussion ────────────────────────────────────────
    _build_discussion(doc, final_report, round_scores)

    # ── 7. Conclusions & Recommendations ─────────────────────
    _build_conclusions(doc, final_report)

    doc.add_page_break()

    # ── 8. References ────────────────────────────────────────
    _build_references(doc, final_report)

    doc.add_page_break()

    # ── Appendix A ───────────────────────────────────────────
    _build_appendix_a(doc, domains, mode, round_scores,
                      total_cost, kur, agent_log, max_rounds)

    doc.add_page_break()

    # ── Appendix B ───────────────────────────────────────────
    _build_appendix_b(doc, agent_log)

    # ── Running headers / footers ─────────────────────────────
    _setup_running_headers(doc, brief_short, doc_no)

    # ── Serialize ─────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# Backward-compat alias
generate_pdf_report = generate_docx_report


# ═══════════════════════════════════════════════════════════════
# TEST
# ═══════════════════════════════════════════════════════════════
if __name__ == "__main__":
    sample_report = """
## 1. EXECUTIVE SUMMARY

Analysis of an AA 6061-T6 aluminum I-beam bridge girder demonstrates adequate structural
performance under the specified loading conditions. The critical finding is that the selected
cross-section (200×100×10 mm I-profile) exceeds the minimum safety factor requirement by a
factor of nearly five, presenting a significant optimization opportunity.

## 2. TECHNICAL FINDINGS BY DOMAIN

### 2.1 Materials Analysis

AA 6061-T6 aluminum alloy was evaluated against the design requirements. Yield strength
is 276 MPa (ASM Handbook Vol. 2). Elastic modulus is 68.9 GPa. Fracture toughness K_IC = 29 MPa√m.

Bending stress analysis:
- Applied UDL: w = 5.18 kN/m (including self-weight of 0.175 kN/m)
- Maximum bending moment: M_max = wL²/8 = 5.18 × 100 / 8 = 64.75 kN·m
- Maximum bending stress: σ_max = M/Z = 27.98 MPa
- Safety factor against yielding: SF = 276 / 27.98 = 9.86

Deflection check:
- Maximum midspan deflection: δ = 5wL⁴/(384EI) = 3.7 mm
- Serviceability ratio: L/δ = 2703 >> L/300 (Eurocode 3 limit)

CROSS-DOMAIN FLAG: Coastal corrosion will reduce web thickness from 8 mm to ~6.75 mm
over 25 years without surface treatment. Structural shear capacity must be re-verified.

## 3. CROSS-DOMAIN ANALYSIS

The materials and structural analyses are consistent. The deflection and stress calculations
use the same cross-sectional properties (I = 2314 cm⁴, Z = 231.4 cm³) and yield compatible
results. The corrosion flag from Materials requires structural re-verification at year 15.

## 4. RISK ASSESSMENT

Failure Mode 1: Corrosion-induced section loss. RPN: 160. Mitigation: anodize coating.
Failure Mode 2: Fatigue cracking at weld toes. RPN: 80. Mitigation: grind smooth.
Failure Mode 3: Lateral-torsional buckling. RPN: 120. Mitigation: add lateral bracing.

## 5. CONCLUSIONS AND RECOMMENDATIONS

CRITICAL: Apply Type II anodize + MIL-PRF-23377 primer before installation.
HIGH: Reduce cross-section to 150×80×8 mm — SF = 5.0 still exceeds target, weight −35%.
HIGH: Verify lateral-torsional buckling per Eurocode 9 §6.3.2 before final design.
MEDIUM: Conduct Miner's Rule fatigue analysis using site-specific load spectrum.

## 6. REFERENCES

ASM Handbook, Volume 2: Properties and Selection (2001).
Eurocode 9: Design of Aluminium Structures (EN 1999-1-1).
MMPDS-11: Metallic Materials Properties Development and Standardization (2016).
"""

    agent_log_test = [
        {"key": "malzeme_a",       "name": "Materials Expert A",    "model": "claude-sonnet-4-6",
         "cost": 0.00742, "output": "Yield strength 276 MPa. Bending stress 27.98 MPa. SF = 9.86.", "thinking": ""},
        {"key": "capraz_dogrulama","name": "Cross-Validation",      "model": "claude-sonnet-4-6",
         "cost": 0.00218, "output": "All calculations dimensionally consistent. I_xx not independently verified.", "thinking": ""},
        {"key": "soru_uretici",    "name": "Question Generator",    "model": "claude-sonnet-4-6",
         "cost": 0.00156, "output": "1. Is end condition truly simply supported? 2. Impact factor for traffic?", "thinking": ""},
        {"key": "gozlemci",        "name": "Observer / Meta-Agent", "model": "claude-sonnet-4-6",
         "cost": 0.00284, "output": "KALİTE PUANI: 87/100. LTB check missing.", "thinking": ""},
        {"key": "final_rapor",     "name": "Final Report Writer",   "model": "claude-opus-4-6",
         "cost": 0.02841, "output": sample_report, "thinking": "Analyzed all domain outputs carefully."},
    ]

    out_bytes = generate_docx_report(
        brief="AA 6061-T6 aluminum bridge girder — 10 m span, 5 kN/m UDL, coastal environment, 25-year design life.",
        final_report=sample_report,
        domains=["Materials", "Structural"],
        round_scores=[{"tur": 1, "puan": 87}],
        agent_log=agent_log_test,
        total_cost=sum(a["cost"] for a in agent_log_test),
        kur=44.0, mode=1,
    )

    out = "/mnt/user-data/outputs/test_v2_report.docx"
    with open(out, "wb") as f:
        f.write(out_bytes)
    print(f"✅  {len(out_bytes):,} bytes  →  {out}")